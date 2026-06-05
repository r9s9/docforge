"""Harder-document tests: optional sections, repeatable sections, and messy
real-world shapes. These intentionally probe the fragile parts of the pipeline.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from docforge.ai_classifier import classify, derive_field_definitions
from docforge.assembler import assemble
from docforge.multi_doc_differ import diff_documents
from docforge.sampledata import build_service_agreement
from docforge.schemas.classification import ClassificationResult, ElementClassification
from docforge.schemas.enums import ClassificationType, FieldType
from docforge.structure_normalizer import build_extraction
from docforge.template_builder import build_template_docx, build_template_from_examples


def _save(doc, path):
    doc.save(str(path))
    return str(path)


# ---------------------------------------------------------------------------
# Optional sections (content present in some examples but not others)
# ---------------------------------------------------------------------------
def test_optional_section_detected_and_toggleable(tmp_path):
    p1 = _save(build_service_agreement(1), tmp_path / "sa1.docx")  # has special clause
    p2 = _save(build_service_agreement(2), tmp_path / "sa2.docx")  # omits it
    template_bytes, _rep, _result, fields = build_template_from_examples([p1, p2])

    include_fields = [f for f in fields if f.field_type == FieldType.BOOLEAN]
    assert include_fields, "optional content should produce an include_* toggle"
    inc = include_fields[0].field_name
    assert inc.startswith("include_")

    base_ctx = {f.field_name: "X" for f in fields if f.field_type == FieldType.TEXT}

    # Included
    on = assemble(template_bytes, {**base_ctx, inc: True}, fields)
    text_on = "\n".join(p.text for p in Document(BytesIO(on)).paragraphs)
    assert "Special Clause" in text_on

    # Excluded
    off = assemble(template_bytes, {**base_ctx, inc: False}, fields)
    text_off = "\n".join(p.text for p in Document(BytesIO(off)).paragraphs)
    assert "Special Clause" not in text_off

    # Default (toggle omitted) -> included (default True)
    default = assemble(template_bytes, base_ctx, fields)
    text_default = "\n".join(p.text for p in Document(BytesIO(default)).paragraphs)
    assert "Special Clause" in text_default


# ---------------------------------------------------------------------------
# Repeatable section (a paragraph rendered once per list item)
# ---------------------------------------------------------------------------
def test_repeatable_section_paragraph(tmp_path):
    doc = Document()
    doc.add_heading("Key Findings", level=1)
    doc.add_paragraph("The first finding goes here.")
    path = _save(doc, tmp_path / "rs.docx")

    ext = build_extraction(path, "rs")
    para_node = next(
        e for e in ext.top_level_elements() if e.type == "paragraph" or e.type.value == "paragraph"
    )
    # pick the body paragraph (not the heading)
    para_node = next(
        e for e in ext.top_level_elements()
        if e.type.value == "paragraph" and "finding" in e.text.lower()
    )

    classifications = []
    for e in ext.elements:
        if e.node_id == para_node.node_id:
            classifications.append(
                ElementClassification(
                    node_id=e.node_id,
                    classification=ClassificationType.REPEATABLE_SECTION,
                    field_name="findings",
                    field_type=FieldType.MULTILINE_TEXT,
                    source="user",
                )
            )
        else:
            classifications.append(
                ElementClassification(
                    node_id=e.node_id, classification=ClassificationType.FIXED, required=False
                )
            )
    result = ClassificationResult(extraction_document_id="rs", classifications=classifications)

    fields = derive_field_definitions(ext, result)
    assert any(f.field_name == "findings" for f in fields)

    template_bytes = build_template_docx(path, result, fields)
    out = assemble(
        template_bytes, {"findings": ["Alpha finding", "Beta finding", "Gamma finding"]}, fields
    )
    texts = [p.text for p in Document(BytesIO(out)).paragraphs if p.text.strip()]
    assert "Alpha finding" in texts
    assert "Beta finding" in texts
    assert "Gamma finding" in texts


# ---------------------------------------------------------------------------
# Messy real-world shapes: multi-section docs, merged cells, info-tables.
# These probe (and honestly document) the fragile parts of the pipeline.
# ---------------------------------------------------------------------------
def test_multisection_document_extracts(tmp_path):
    doc = Document()
    doc.add_heading("Part One", level=1)
    doc.add_paragraph("Reference: A-1")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Part Two", level=1)
    doc.add_paragraph("Reference: A-2")
    path = _save(doc, tmp_path / "multi.docx")

    ext = build_extraction(path, "multi")
    # Multiple sections detected; extraction does not crash; headings present.
    assert len(ext.sections) >= 2
    assert any(e.type.value == "heading" for e in ext.elements)


def test_merged_cell_table_extracts(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    # Merge the two header cells into one spanning cell.
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Summary"
    table.cell(1, 0).text = "Owner"
    table.cell(1, 1).text = "Jane"
    table.cell(2, 0).text = "Status"
    table.cell(2, 1).text = "Open"
    path = _save(doc, tmp_path / "merged.docx")

    ext = build_extraction(path, "merged")
    tables = [e for e in ext.elements if e.type.value == "table"]
    assert tables, "merged-cell table should still be detected"
    ts = tables[0].table_structure
    assert ts is not None and ts.merged_cells, "the horizontal merge should be recorded"


def _infotable_doc(variant: int) -> Document:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ACCOUNT SUMMARY").bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    pairs = [("Client", f"Acme {variant}"), ("Date", f"2026-0{variant}-01"), ("Balance", f"${variant}00")]
    for i, (k, v) in enumerate(pairs):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    return doc


@pytest.mark.xfail(
    reason="KNOWN GAP: fixed-layout label/value tables aren't split into per-cell "
    "dynamic fields (whole table is treated as fixed/repeatable).",
    strict=False,
)
def test_infotable_percell_dynamic_fields(tmp_path):
    p1 = _save(_infotable_doc(1), tmp_path / "info1.docx")
    p2 = _save(_infotable_doc(2), tmp_path / "info2.docx")
    exts = [build_extraction(p1, "i1"), build_extraction(p2, "i2")]
    diff = diff_documents(exts)
    result = classify(exts[0], diff)
    fields = {f.field_name for f in derive_field_definitions(exts[0], result)}
    # Ideal behaviour: the changing cells become their own dynamic fields.
    assert {"client", "date", "balance"} & fields
