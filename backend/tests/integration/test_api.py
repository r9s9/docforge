"""End-to-end API tests via FastAPI TestClient (spec §15 endpoints)."""

from __future__ import annotations

import jwt
import pytest
from fastapi.testclient import TestClient

from docforge.api.app import app
from docforge.api.auth import CurrentUser, get_current_user
from docforge.api.deps import get_db

DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Fixed identity for the behavioral tests (the auth-specific tests below use real
# signed JWTs instead).
TEST_USER_ID = "00000000-0000-0000-0000-000000000001"
TEST_JWT_SECRET = "test-jwt-secret-at-least-32-bytes-long-xxxx"


@pytest.fixture
def client(db_session, settings_tmp):
    """TestClient with the DB dependency pointed at the isolated test session.

    Instantiated without the lifespan context so the module-level engine/init_db
    is never touched; file storage is redirected by the settings_tmp fixture.
    Authentication is stubbed to a fixed user so these tests focus on behavior;
    the auth tests further down verify token handling and per-user isolation.
    """

    def _override_db():
        yield db_session

    app.dependency_overrides[get_db] = _override_db
    app.dependency_overrides[get_current_user] = lambda: CurrentUser(
        id=TEST_USER_ID, email="tester@example.com"
    )
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.clear()


def _make_token(secret: str, user_id: str, email: str = "user@example.com") -> str:
    return jwt.encode(
        {"sub": user_id, "email": email, "aud": "authenticated"}, secret, algorithm="HS256"
    )


@pytest.fixture
def auth_client(db_session, settings_tmp, monkeypatch):
    """TestClient that runs the *real* JWT verification (no auth override).

    Configures the test JWT secret so tokens minted with ``_make_token`` verify.
    """
    monkeypatch.setattr(settings_tmp, "auth_required", True)
    monkeypatch.setattr(settings_tmp, "supabase_jwt_secret", TEST_JWT_SECRET)
    monkeypatch.setattr(settings_tmp, "supabase_jwt_audience", "authenticated")

    def _override_db():
        yield db_session

    app.dependency_overrides[get_db] = _override_db
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.clear()


def _upload(client, docs):
    files = [("files", (p.name, p.read_bytes(), DOCX)) for p in docs]
    return client.post("/api/templates/analyze", files=files)


def _publish_via_services(db, settings, docs, name="Template", owner_id=TEST_USER_ID):
    """Build a published template directly via the service layer.

    Analysis runs in a background thread that uses the *real* engine, so it
    can't see the TestClient's isolated session — we exercise the HTTP lifecycle
    endpoints against a template created synchronously here instead.
    """
    from pathlib import Path

    from docforge.document_ingest import store_source_document
    from docforge.services import analyze_documents, publish_template
    from docforge.template_registry import TemplateRegistry

    sources = [
        store_source_document(db, p.name, Path(p).read_bytes(), owner_id=owner_id) for p in docs
    ]
    job = analyze_documents(db, sources, settings=settings)
    registry = TemplateRegistry(settings.templates_dir)
    template, _ = publish_template(
        db, job, name=name, settings=settings, registry=registry, owner_id=owner_id
    )
    return template.id


def test_health(client):
    r = client.get("/api/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"


def test_analyze_accepts_and_queues(client, project_docs):
    r = _upload(client, project_docs)
    assert r.status_code == 202, r.text
    assert r.json()["status"] in ("pending", "running", "completed")


def test_full_lifecycle(client, db_session, settings_tmp, project_docs):
    template_id = _publish_via_services(db_session, settings_tmp, project_docs, "Project Report")

    # browse
    assert any(t["id"] == template_id for t in client.get("/api/templates").json())
    detail = client.get(f"/api/templates/{template_id}").json()
    assert detail["latest"]["fields"]
    versions = client.get(f"/api/templates/{template_id}/versions").json()
    assert versions[0]["version"] == 1

    # 4) generate (structured)
    body = {
        "mode": "structured_json",
        "data": {
            "project_name": "Orion",
            "report_date": "2026-07-01",
            "prepared_by": "Alice Brown",
            "summary": "On track.",
            "task_status": [
                {"task": "Design", "owner": "M. Lee", "status": "Done", "due_date": "2026-07-01"},
            ],
        },
    }
    r = client.post(f"/api/templates/{template_id}/generate", json=body)
    assert r.status_code == 200, r.text
    gen = r.json()
    assert gen["validation"]["status"] == "pass"

    # 5) download the generated docx
    dl = client.get(gen["download_url"])
    assert dl.status_code == 200
    assert len(dl.content) > 1000
    assert "wordprocessingml" in dl.headers["content-type"]

    # 6) download the template docx
    tdl = client.get(f"/api/templates/{template_id}/versions/1/template.docx")
    assert tdl.status_code == 200


def test_route_and_validate(client, db_session, settings_tmp, project_docs):
    template_id = _publish_via_services(db_session, settings_tmp, project_docs, "PR")

    # route preview (unstructured)
    r = client.post(
        f"/api/templates/{template_id}/route",
        json={"raw_text": "Project Name: Helios\nReport Date: 2026-08-01\nGoing well."},
    )
    assert r.status_code == 200
    placements = {p["field_name"]: p["value"] for p in r.json()["placements"]}
    assert placements.get("project_name") == "Helios"

    # validate a deliberately incomplete context -> fail
    r = client.post(f"/api/templates/{template_id}/validate", json={"context": {"project_name": "X"}})
    assert r.status_code == 200
    assert r.json()["status"] in ("fail", "warning")


def test_rename_and_delete_template(client, db_session, settings_tmp, project_docs):
    template_id = _publish_via_services(db_session, settings_tmp, project_docs, "Original Name")

    # rename
    r = client.patch(f"/api/templates/{template_id}", json={"name": "Renamed Template"})
    assert r.status_code == 200
    assert r.json()["name"] == "Renamed Template"

    # delete
    r = client.delete(f"/api/templates/{template_id}")
    assert r.status_code == 204
    assert client.get(f"/api/templates/{template_id}").status_code == 404
    assert all(t["id"] != template_id for t in client.get("/api/templates").json())


def test_analyze_rejects_too_many_files(client, project_docs, monkeypatch, settings_tmp):
    monkeypatch.setattr(settings_tmp, "max_files_per_analysis", 1)
    r = _upload(client, project_docs)  # 2 files, limit 1
    assert r.status_code == 400


def _make_job(db, settings, docs, owner_id=TEST_USER_ID):
    from pathlib import Path

    from docforge.document_ingest import store_source_document

    from docforge.services import analyze_documents

    sources = [
        store_source_document(db, p.name, Path(p).read_bytes(), owner_id=owner_id) for p in docs
    ]
    job = analyze_documents(db, sources, settings=settings)
    job.owner_id = owner_id
    db.commit()
    return job


def test_analysis_preview_docx_filled_and_tags(client, db_session, settings_tmp, project_docs):
    import io

    from docx import Document

    job = _make_job(db_session, settings_tmp, project_docs)

    # Sample-filled: a real document with readable «Label» tokens, no raw Jinja.
    r = client.post(f"/api/analyses/{job.id}/preview.docx?mode=filled")
    assert r.status_code == 200, r.text
    assert "wordprocessingml" in r.headers["content-type"]
    Document(io.BytesIO(r.content))  # opens as a valid DOCX
    filled_text = "\n".join(p.text for p in Document(io.BytesIO(r.content)).paragraphs)
    assert "«" in filled_text
    assert "{{" not in filled_text

    # Template tags: the raw template with visible Jinja placeholders.
    r2 = client.post(f"/api/analyses/{job.id}/preview.docx?mode=tags")
    assert r2.status_code == 200, r2.text
    tags_text = "\n".join(p.text for p in Document(io.BytesIO(r2.content)).paragraphs)
    assert "{{" in tags_text


def test_analysis_preview_docx_unknown_job_404(client):
    r = client.post("/api/analyses/does-not-exist/preview.docx?mode=filled")
    assert r.status_code == 404


def _edit_fixed_text(path):
    """Return DOCX bytes of project_docs[0] with a FIXED heading altered."""
    import io

    from docx import Document

    doc = Document(str(path))
    doc.paragraphs[0].text = doc.paragraphs[0].text + " [EDITED]"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_compliance_alignment_and_fix(client, db_session, settings_tmp, project_docs):
    import io

    from docx import Document

    template_id = _publish_via_services(db_session, settings_tmp, project_docs, "Compliance T")
    candidate = _edit_fixed_text(project_docs[0])

    # 1) Compliance check returns side-by-side alignment + a fixable flag.
    r = client.post(
        f"/api/templates/{template_id}/compliance",
        files=[("file", ("cand.docx", candidate, DOCX))],
    )
    assert r.status_code == 200, r.text
    report = r.json()
    assert report["alignment"], "alignment data should be present"
    assert report["fixable"] is True
    assert any(p["status"] == "changed" for p in report["alignment"])

    # 2) Fix endpoint restores the boilerplate and reports the count.
    rf = client.post(
        f"/api/templates/{template_id}/compliance/fix",
        files=[("file", ("cand.docx", candidate, DOCX))],
    )
    assert rf.status_code == 200, rf.text
    assert "wordprocessingml" in rf.headers["content-type"]
    assert int(rf.headers.get("X-Fixes-Applied", "0")) >= 1
    body = "\n".join(p.text for p in Document(io.BytesIO(rf.content)).paragraphs)
    assert "[EDITED]" not in body  # the altered boilerplate was restored


# --------------------------------------------------------------------------- #
# Authentication & per-user isolation                                          #
# --------------------------------------------------------------------------- #

USER_A = "11111111-1111-1111-1111-111111111111"
USER_B = "22222222-2222-2222-2222-222222222222"


def test_requires_token(auth_client):
    """Data routes reject requests with no/invalid bearer token."""
    assert auth_client.get("/api/templates").status_code == 401
    assert auth_client.get("/api/me").status_code == 401
    bad = {"Authorization": "Bearer not-a-real-token"}
    assert auth_client.get("/api/templates", headers=bad).status_code == 401


def test_me_returns_identity(auth_client):
    tok = _make_token(TEST_JWT_SECRET, USER_A, "a@example.com")
    r = auth_client.get("/api/me", headers={"Authorization": f"Bearer {tok}"})
    assert r.status_code == 200, r.text
    assert r.json() == {"id": USER_A, "email": "a@example.com"}


def test_token_signed_with_wrong_secret_rejected(auth_client):
    tok = _make_token("the-wrong-secret", USER_A)
    r = auth_client.get("/api/templates", headers={"Authorization": f"Bearer {tok}"})
    assert r.status_code == 401


def test_templates_scoped_per_user(auth_client, db_session, settings_tmp, project_docs):
    """User A's template is visible only to A; B sees an empty list and 404s."""
    tid = _publish_via_services(db_session, settings_tmp, project_docs, "A's Template", owner_id=USER_A)
    a = {"Authorization": f"Bearer {_make_token(TEST_JWT_SECRET, USER_A)}"}
    b = {"Authorization": f"Bearer {_make_token(TEST_JWT_SECRET, USER_B)}"}

    # A sees and can open its template.
    a_list = auth_client.get("/api/templates", headers=a).json()
    assert [t["id"] for t in a_list] == [tid]
    assert auth_client.get(f"/api/templates/{tid}", headers=a).status_code == 200

    # B's list is empty and B cannot open or download A's template (404, no leak).
    assert auth_client.get("/api/templates", headers=b).json() == []
    assert auth_client.get(f"/api/templates/{tid}", headers=b).status_code == 404
    dl = f"/api/templates/{tid}/versions/1/template.docx"
    assert auth_client.get(dl, headers=a).status_code == 200
    assert auth_client.get(dl, headers=b).status_code == 404
