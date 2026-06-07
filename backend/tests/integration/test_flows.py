"""Integration tests for the three end-to-end user flows (spec §5)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from docforge.document_ingest import store_source_document
from docforge.schemas.enums import GenerationMode, JobStatus, ValidationStatus
from docforge.schemas.generation import GenerationInput
from docforge.services import analyze_documents, generate_document, publish_template
from docforge.storage import get_storage
from docforge.template_registry import TemplateRegistry


def _output_doc(gen) -> Document:
    """Load a generated document from storage (output_path is a storage key)."""
    assert gen.output_path and get_storage().exists(gen.output_path)
    return Document(BytesIO(get_storage().get_bytes(gen.output_path)))


def _ingest(db, docs):
    return [store_source_document(db, p.name, Path(p).read_bytes()) for p in docs]


def _publish(db, settings, docs, name):
    sources = _ingest(db, docs)
    job = analyze_documents(db, sources, settings=settings)
    registry = TemplateRegistry(settings.templates_dir)
    template, version = publish_template(db, job, name=name, settings=settings, registry=registry)
    return job, template, version, registry


def test_flow1_analyze_and_publish_creates_package(db_session, settings_tmp, project_docs):
    job, template, version, registry = _publish(db_session, settings_tmp, project_docs, "Project Report")
    assert job.status == JobStatus.COMPLETED.value
    assert job.field_definitions
    assert version.version == 1

    assert registry.version_exists(template.id, 1)
    base = f"templates/{template.id}/1"
    for artifact in (
        "template.docx",
        "manifest.json",
        "field_definitions.json",
        "validation_rules.json",
        "template_intelligence.json",
        "review_snapshot.json",
    ):
        assert registry.storage.exists(f"{base}/{artifact}"), f"missing {artifact}"
    assert registry.source_example_names(template.id, 1)  # source examples saved


def test_flow2_structured_generation(db_session, settings_tmp, project_docs):
    _, template, _, registry = _publish(db_session, settings_tmp, project_docs, "Project Report")
    gen_input = GenerationInput(
        mode=GenerationMode.STRUCTURED_JSON,
        data={
            "project_name": "Orion Platform",
            "report_date": "2026-07-01",
            "prepared_by": "Alice Brown",
            "summary": "All milestones on track.",
            "task_status": [
                {"task": "Design", "owner": "M. Lee", "status": "Done", "due_date": "2026-07-01"},
                {"task": "Build", "owner": "P. Ono", "status": "WIP", "due_date": "2026-07-20"},
            ],
        },
    )
    gen = generate_document(db_session, template, gen_input, settings=settings_tmp, registry=registry)
    assert gen.validation["status"] == ValidationStatus.PASS.value

    doc = _output_doc(gen)
    assert any("Project Name: Orion Platform" in p.text for p in doc.paragraphs)
    assert len(doc.tables[0].rows) == 3  # header + 2 rows


def test_flow3_unstructured_generation(db_session, settings_tmp, project_docs):
    _, template, _, registry = _publish(db_session, settings_tmp, project_docs, "Project Report")
    raw = (
        "Project Name: Helios CRM\n"
        "Report Date: 2026-08-01\n"
        "Prepared By: Sam Lee\n"
        "Everything is on schedule and the pilot is complete."
    )
    gen_input = GenerationInput(mode=GenerationMode.UNSTRUCTURED_TEXT, raw_text=raw)
    gen = generate_document(db_session, template, gen_input, settings=settings_tmp, registry=registry)
    doc = _output_doc(gen)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Project Name: Helios CRM" in body
    assert "Report Date: 2026-08-01" in body


def test_invoice_end_to_end(db_session, settings_tmp, invoice_docs):
    _, template, _, registry = _publish(db_session, settings_tmp, invoice_docs, "Invoice")
    gen_input = GenerationInput(
        mode=GenerationMode.STRUCTURED_JSON,
        data={
            "invoice_number": "INV-9000",
            "invoice_date": "2026-09-01",
            "bill_to": "Initech",
            "total_due": "$4,200.00",
            "line_items": [
                {"description": "Retainer", "qty": "1", "unit_price": "$4,200.00", "amount": "$4,200.00"},
            ],
        },
    )
    gen = generate_document(db_session, template, gen_input, settings=settings_tmp, registry=registry)
    assert get_storage().exists(gen.output_path)
