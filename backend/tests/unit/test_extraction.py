"""Unit tests for OOXML extraction + normalization."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docforge.ooxml_extractor import DocxPackage, read_raw_parts
from docforge.schemas.enums import ElementType
from docforge.structure_normalizer import build_extraction


def _fldchar(kind: str):
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), kind)
    return fc


def _instr_text(text: str):
    it = OxmlElement("w:instrText")
    it.text = text
    return it


def test_extraction_finds_core_elements(project_docs):
    ext = build_extraction(project_docs[0], "doc-1")
    texts = [e.text for e in ext.elements]
    assert "MONTHLY PROJECT STATUS REPORT" in texts
    # labeled dynamic line present
    assert any(t.startswith("Project Name:") for t in texts)
    # at least one heading
    assert any(e.type == ElementType.HEADING for e in ext.elements)


def test_extraction_table_structure(project_docs):
    ext = build_extraction(project_docs[0], "doc-1")
    tables = [e for e in ext.elements if e.type == ElementType.TABLE]
    assert len(tables) == 1
    ts = tables[0].table_structure
    assert ts is not None
    assert ts.headers == ["Task", "Owner", "Status", "Due Date"]
    # header row + 2 data rows in variant 1
    assert ts.n_rows == 3
    assert ts.n_cols == 4


def test_extraction_detects_footer_auto_field(project_docs):
    ext = build_extraction(project_docs[0], "doc-1")
    footer_nodes = [e for e in ext.elements if e.header_footer_scope and "footer" in e.header_footer_scope]
    assert footer_nodes, "expected at least one footer element"
    assert any("auto_field" in e.semantic_hints for e in footer_nodes)


def test_node_ids_are_stable(project_docs):
    a = build_extraction(project_docs[0], "doc-1")
    b = build_extraction(project_docs[0], "doc-1")
    assert [e.node_id for e in a.elements] == [e.node_id for e in b.elements]
    assert [(e.type, e.text) for e in a.elements] == [(e.type, e.text) for e in b.elements]


def test_content_hash_differs_between_variants(project_docs):
    a = build_extraction(project_docs[0], "doc-1")
    b = build_extraction(project_docs[1], "doc-2")
    assert a.content_hash != b.content_hash


def test_raw_parts_inventory(project_docs):
    pkg = DocxPackage.from_path(project_docs[0])
    info = read_raw_parts(pkg)
    assert info["main_document"] == "word/document.xml"
    assert info["has_styles"] is True
    assert info["n_parts"] > 3


def test_extraction_finds_text_wrapped_in_content_control(tmp_path):
    """A paragraph wrapped in a <w:sdt> (Rich Text / Plain Text content control)
    used to be completely invisible to extraction — python-docx has no typed
    wrapper for <w:sdt>, so a plain isinstance check silently skipped it."""
    doc = Document()
    doc.add_paragraph("Plain paragraph before.")
    wrapped = doc.add_paragraph("Text inside a content control.")
    p_elm = wrapped._p
    body = p_elm.getparent()
    sdt = OxmlElement("w:sdt")
    sdt.append(OxmlElement("w:sdtPr"))
    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)
    body.replace(p_elm, sdt)
    sdt_content.append(p_elm)
    doc.add_paragraph("Plain paragraph after.")

    path = tmp_path / "sdt.docx"
    doc.save(str(path))

    ext = build_extraction(str(path), "sdt-doc")
    texts = [e.text for e in ext.top_level_elements()]
    assert texts == [
        "Plain paragraph before.",
        "Text inside a content control.",
        "Plain paragraph after.",
    ]


def test_extraction_toc_field_spans_multiple_paragraphs(tmp_path):
    """A generated Table of Contents puts the fldChar begin/instrText marker
    only in the FIRST paragraph — every visible entry after that ("1.1 Agenda
    & Topics ... 3") is its own plain paragraph with no field marker of its
    own, until a final `end` marker closes the field several paragraphs
    later. Checking each paragraph in isolation used to only recognize that
    first line as TOC/auto-field content, leaving every other entry looking
    like ordinary fixed text (eligible to become a broken static placeholder
    in tags-only mode, when it should stay a live, self-updating field)."""
    doc = Document()

    p1 = doc.add_paragraph()
    p1.add_run()._element.append(_fldchar("begin"))
    p1.add_run()._element.append(_instr_text(' TOC \\o "1-3" \\h \\z \\u '))
    p1.add_run()._element.append(_fldchar("separate"))
    p1.add_run("Heading One\t1")

    doc.add_paragraph("Heading Two\t2")  # no field marker of its own

    p3 = doc.add_paragraph("Heading Three\t3")
    p3.add_run()._element.append(_fldchar("end"))

    doc.add_paragraph("Plain paragraph after the TOC.")

    path = tmp_path / "toc.docx"
    doc.save(str(path))

    ext = build_extraction(str(path), "toc-doc")
    toc_paras = [e for e in ext.top_level_elements() if e.text.startswith("Heading")]
    assert len(toc_paras) == 3
    for e in toc_paras:
        assert "auto_field" in e.semantic_hints, f"{e.text!r} missing auto_field hint"
        assert "toc" in e.semantic_hints, f"{e.text!r} missing toc hint"

    after = next(e for e in ext.top_level_elements() if "after" in e.text)
    assert "toc" not in after.semantic_hints  # field closed correctly — doesn't leak
    assert "auto_field" not in after.semantic_hints
