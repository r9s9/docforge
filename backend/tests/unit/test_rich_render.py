"""Rich (markdown-lite) field values: parsing and rendering into real paragraphs."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docforge.assembler import assemble
from docforge.assembler.richtext import (
    is_rich,
    parse_rich_blocks,
    parse_spans,
    strip_markers,
)
from docforge.schemas.enums import ClassificationType, FieldType
from docforge.schemas.template import FieldDefinition


# --- parsing ----------------------------------------------------------------


def test_is_rich_only_for_structured_values():
    assert is_rich("line one\nline two")
    assert is_rich("- a bullet")
    assert is_rich("has **bold** in it")
    assert not is_rich("Acme Corporation")
    assert not is_rich("")
    # Lone asterisks in prose are not emphasis.
    assert not is_rich("5 * 4 = 20")
    assert not is_rich("Terms apply*")


def test_parse_rich_blocks_paragraphs_bullets_and_numbers():
    blocks = parse_rich_blocks(
        "Opening line.\n\nSecond paragraph.\n- first\n- second\n1. step one\n2. step two"
    )
    assert [b.kind for b in blocks] == [
        "paragraph",
        "paragraph",
        "bullet",
        "bullet",
        "number",
        "number",
    ]
    assert [b.text for b in blocks][:2] == ["Opening line.", "Second paragraph."]
    # Numbered markers keep the author's own numbers.
    assert [b.marker for b in blocks[-2:]] == ["1.", "2."]


def test_single_newlines_separate_paragraphs():
    # A model that separates paragraphs with one newline means the same thing a
    # user does — joining them back is the bug this format exists to fix.
    assert len(parse_rich_blocks("one\ntwo\nthree")) == 3


def test_parse_spans_extracts_emphasis_and_drops_markers():
    spans = parse_spans("plain **bold** and *italic* end")
    assert [(s.text, s.bold, s.italic) for s in spans] == [
        ("plain ", False, False),
        ("bold", True, False),
        (" and ", False, False),
        ("italic", False, True),
        (" end", False, False),
    ]


def test_nested_bullets_get_levels():
    blocks = parse_rich_blocks("- top\n  - nested\n    - deeper")
    assert [b.level for b in blocks] == [0, 1, 2]


def test_strip_markers_returns_plain_text():
    assert strip_markers("- a **bold** item") == "a bold item"
    assert strip_markers("2. numbered") == "numbered"


# --- rendering --------------------------------------------------------------


def _template(placeholder: str = "{{ summary }}", *, prefix: str = "") -> bytes:
    """A one-paragraph template whose value run carries distinctive formatting."""
    doc = Document()
    para = doc.add_paragraph()
    if prefix:
        para.add_run(prefix).bold = True
    run = para.add_run(placeholder)
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x22, 0x44, 0x88)
    doc.add_paragraph("Fixed boilerplate.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _field(name: str = "summary", **kw) -> FieldDefinition:
    kw.setdefault("field_type", FieldType.MULTILINE_TEXT)
    kw.setdefault("classification", ClassificationType.DYNAMIC_TEXT)
    return FieldDefinition(field_name=name, **kw)


def _rendered(value: str, *, prefix: str = "") -> Document:
    out = assemble(_template(prefix=prefix), {"summary": value}, [_field()])
    return Document(BytesIO(out))


def test_multiple_paragraphs_become_separate_paragraphs():
    doc = _rendered("First para.\n\nSecond para.\n\nThird para.")
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["First para.", "Second para.", "Third para.", "Fixed boilerplate."]


def test_generated_paragraphs_inherit_the_placeholder_formatting():
    doc = _rendered("One.\n\nTwo.")
    second = doc.paragraphs[1]
    run = second.runs[0]
    assert run.font.size == Pt(13)
    assert run.italic is True
    assert run.font.color.rgb == RGBColor(0x22, 0x44, 0x88)


def test_emphasis_is_added_without_dropping_inherited_styling():
    doc = _rendered("plain **strong** words")
    runs = doc.paragraphs[0].runs
    strong = next(r for r in runs if r.text == "strong")
    assert strong.bold is True
    # The template's own italic + size survive on the emphasised run.
    assert strong.italic is True
    assert strong.font.size == Pt(13)


def test_static_prefix_keeps_the_first_block_on_its_line():
    doc = _rendered("Alpha.\n\nBeta.", prefix="Summary: ")
    assert doc.paragraphs[0].text == "Summary: Alpha."
    assert doc.paragraphs[1].text == "Beta."


def test_bullets_use_the_documents_own_numbering_when_available():
    doc = _rendered("- alpha\n- beta")
    for para in doc.paragraphs[:2]:
        pPr = para._p.find(qn("w:pPr"))
        assert pPr is not None and pPr.find(qn("w:numPr")) is not None


def test_numbered_items_keep_literal_markers_and_indent():
    # Word numbering is deliberately not reused for ordered lists: a borrowed
    # definition carries a running counter and would renumber the list.
    doc = _rendered("1. first\n2. second")
    assert [p.text for p in doc.paragraphs[:2]] == ["1. first", "2. second"]
    assert doc.paragraphs[0].paragraph_format.left_indent is not None


def test_bullets_fall_back_to_literal_markers_without_a_numbering_part():
    doc = Document()
    doc.add_paragraph().add_run("{{ summary }}")
    buf = BytesIO()
    doc.save(buf)
    stripped = _without_numbering(buf.getvalue())
    out = assemble(stripped, {"summary": "- alpha\n- beta"}, [_field()])
    rendered = Document(BytesIO(out))
    assert [p.text for p in rendered.paragraphs[:2]] == ["• alpha", "• beta"]


def _without_numbering(docx_bytes: bytes) -> bytes:
    """Rebuild the package without numbering.xml (some templates have none).

    The part, its relationship and its content-type override all have to go
    together, or the package no longer opens.
    """
    import re
    import zipfile

    src = zipfile.ZipFile(BytesIO(docx_bytes))
    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename.endswith("numbering.xml"):
                continue
            data = src.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = re.sub(rb"<Override[^>]*numbering\.xml[^>]*/>", b"", data)
            elif item.filename.endswith("document.xml.rels"):
                data = re.sub(rb"<Relationship[^>]*numbering\.xml[^>]*/>", b"", data)
            dst.writestr(item, data)
    return out.getvalue()


def test_rich_values_expand_inside_table_cells():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("{{ summary }}")
    buf = BytesIO()
    doc.save(buf)
    out = assemble(buf.getvalue(), {"summary": "cell one\n\ncell two"}, [_field()])
    cell = Document(BytesIO(out)).tables[0].cell(0, 0)
    assert [p.text for p in cell.paragraphs] == ["cell one", "cell two"]


def test_plain_values_are_untouched_by_the_rich_path():
    doc = _rendered("Just one ordinary line.")
    assert [p.text for p in doc.paragraphs] == ["Just one ordinary line.", "Fixed boilerplate."]
    assert len(doc.paragraphs[0].runs) == 1


def test_scalar_field_types_never_go_through_rich_rendering():
    # A date that somehow contains a newline must not sprout paragraphs.
    field = _field("when", field_type=FieldType.DATE)
    doc = Document()
    doc.add_paragraph().add_run("{{ when }}")
    buf = BytesIO()
    doc.save(buf)
    out = assemble(buf.getvalue(), {"when": "2026-01-01\n2026-02-02"}, [field])
    assert len(Document(BytesIO(out)).paragraphs) == 1


def test_repeatable_sections_strip_markdown_markers():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("{%p for item in points %}")
    doc.add_paragraph().add_run("{{ item }}")
    doc.add_paragraph().add_run("{%p endfor %}")
    buf = BytesIO()
    doc.save(buf)
    field = _field("points", classification=ClassificationType.REPEATABLE_SECTION)
    out = assemble(buf.getvalue(), {"points": "- alpha\n- **beta**"}, [field])
    texts = [p.text for p in Document(BytesIO(out)).paragraphs if p.text.strip()]
    assert texts == ["alpha", "beta"]
