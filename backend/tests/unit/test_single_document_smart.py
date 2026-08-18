"""Smart detect with one example: content is content, not boilerplate.

Smart mode decides fixed-vs-dynamic by comparing example documents. Given only
one there is nothing to compare, and "no evidence of change" was being read as
"this never changes" — so a document's whole body came through fixed and the
template could not be filled in at all.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from docforge.ai_classifier import classify, derive_field_definitions
from docforge.assembler import assemble
from docforge.schemas.enums import ClassificationType
from docforge.structure_normalizer import build_extraction
from docforge.template_builder.builder import build_template_docx


@pytest.fixture
def training_doc(tmp_path):
    """A concept document: headings, a metadata block, and bullet lists."""
    doc = Document()
    doc.add_heading("Copilot Agent Builder Training", level=1)
    doc.add_paragraph("Target Audience: Experienced M365 users")
    doc.add_paragraph("Duration: 90 minutes")
    doc.add_heading("Welcome & Objectives", level=2)
    for line in ("Quick recap: what is Copilot?", "What is an agent?", "What we will build"):
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Copilot Studio Overview", level=2)
    for line in ("Key components: topics, triggers, actions", "Agent lifecycle", "Environments"):
        doc.add_paragraph(line, style="List Bullet")
    src = tmp_path / "training.docx"
    doc.save(src)
    return str(src)


def _analyze(src: str):
    extraction = build_extraction(src, "d0")
    result = classify(extraction, None, mode="smart")  # one example, no diff
    return extraction, result, derive_field_definitions(extraction, result)


def test_body_content_becomes_fillable(training_doc):
    extraction, result, fields = _analyze(training_doc)
    by_node = {c.node_id: c for c in result.classifications}
    bullets = [
        e for e in extraction.top_level_elements()
        if (e.text or "").strip().startswith(("Quick recap", "What is an agent", "Key components"))
    ]
    assert bullets, "fixture produced no bullets"
    for e in bullets:
        assert by_node[e.node_id].classification != ClassificationType.FIXED, (
            f"{e.text[:40]!r} stayed fixed — the template could not be filled in"
        )


def test_a_bullet_list_is_one_field_not_one_per_line(training_doc):
    """Forty bullets must not become forty cards to scroll past."""
    _extraction, result, _fields = _analyze(training_doc)
    runs = [c for c in result.classifications if c.classification == ClassificationType.REPEATABLE_SECTION]
    assert runs, "consecutive body lines were not grouped"
    assert len({c.field_name for c in runs}) < len(runs)


def test_headings_and_labels_stay_boilerplate(training_doc):
    """Smart detect's promise: the skeleton stays put."""
    extraction, result, _fields = _analyze(training_doc)
    by_node = {c.node_id: c for c in result.classifications}
    headings = [e for e in extraction.top_level_elements() if e.type.value == "heading"]
    assert headings
    for e in headings:
        assert by_node[e.node_id].classification == ClassificationType.FIXED


def test_labelled_metadata_still_becomes_its_own_field(training_doc):
    """"Duration: 90 minutes" is a value with a label, not prose."""
    _extraction, _result, fields = _analyze(training_doc)
    names = {f.field_name for f in fields}
    assert "target_audience" in names and "duration" in names


def test_the_built_template_compiles_and_renders(training_doc):
    """A section toggle must never be closed across a loop — that is invalid Jinja
    and fails the whole template, not just the toggle."""
    _extraction, result, fields = _analyze(training_doc)
    template = build_template_docx(training_doc, result, fields)
    section = next(
        f.field_name for f in fields
        if f.classification == ClassificationType.REPEATABLE_SECTION
    )
    rendered = Document(BytesIO(assemble(template, {section: ["ONE", "TWO"]}, fields)))
    text = [p.text for p in rendered.paragraphs]
    assert "ONE" in text and "TWO" in text
