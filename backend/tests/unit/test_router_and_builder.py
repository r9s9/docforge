"""Unit tests for content routing and template build+assemble."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docforge.ai_router import route_structured, route_unstructured_heuristic
from docforge.assembler import assemble
from docforge.assembler.assembler import _image_bytes, build_render_context
from docforge.schemas.enums import FieldType
from docforge.schemas.template import FieldDefinition
from docforge.template_builder import build_template_from_examples
from docforge.template_builder.builder import (
    _PIC_NS,
    _neutralize_run,
    _neutralize_stray_tags,
    _paragraph_has_picture,
    _run_has_image,
    _tag_picture,
    _templatize_paragraph,
    _templatize_table,
)


def _fields():
    return [
        FieldDefinition(field_name="project_name", label="Project Name", field_type=FieldType.TEXT, required=True),
        FieldDefinition(field_name="report_date", label="Report Date", field_type=FieldType.DATE, required=True),
        FieldDefinition(field_name="summary", label="Summary", field_type=FieldType.MULTILINE_TEXT, required=False),
    ]


def test_route_structured_reports_missing_required():
    res = route_structured(_fields(), {"project_name": "X"}, "t", 1)
    assert "report_date" in res.missing_required
    assert {p.field_name for p in res.placements} == {"project_name"}


def test_route_unstructured_matches_labels_and_prose():
    text = "Project Name: Orion\nReport Date: 2026-07-01\nWe finished the pilot successfully."
    res = route_unstructured_heuristic(_fields(), text, "t", 1)
    by = {p.field_name: p.value for p in res.placements}
    assert by["project_name"] == "Orion"
    assert by["report_date"] == "2026-07-01"
    assert "summary" in by  # leftover prose routed to the free-text field


def test_route_unstructured_distributes_prose_across_many_fields():
    """A tags-only template has MANY text fields — leftover prose should spread
    across whichever fields it topically matches, not dump into the first one."""
    fields = [
        FieldDefinition(
            field_name="session_goals", label="Session Goals", field_type=FieldType.MULTILINE_TEXT,
            description="The training session's learning goals and objectives.", required=True,
        ),
        FieldDefinition(
            field_name="logistics_notes", label="Logistics Notes", field_type=FieldType.MULTILINE_TEXT,
            description="Room booking, projector and equipment logistics for the session.", required=True,
        ),
    ]
    text = (
        "Our main learning goals and objectives this quarter are to improve onboarding.\n"
        "We booked the room and projector equipment for logistics ahead of time."
    )
    res = route_unstructured_heuristic(fields, text, "t", 1)
    by = {p.field_name: p.value for p in res.placements}
    assert "goals" in by["session_goals"].lower()
    assert "projector" in by["logistics_notes"].lower()
    # Neither field absorbed the other's sentence.
    assert "projector" not in by["session_goals"].lower()
    assert "goals" not in by["logistics_notes"].lower()


def test_route_unstructured_fallback_prefers_multiline_over_title_field():
    """When nothing clears the match threshold (e.g. generic tags-only-mode
    descriptions with AI off), the last-resort dump should prefer a
    multiline/body field over a short title-style TEXT field."""
    fields = [
        FieldDefinition(
            field_name="section_title", label="Section Title", field_type=FieldType.TEXT,
            description="Content for this section.", required=True,
        ),
        FieldDefinition(
            field_name="section_body", label="Section Body", field_type=FieldType.MULTILINE_TEXT,
            description="Content for this section.", required=True,
        ),
    ]
    text = "Some unrelated prose that shares no keywords with either field description."
    res = route_unstructured_heuristic(fields, text, "t", 1)
    by = {p.field_name: p.value for p in res.placements}
    assert "section_body" in by
    assert "section_title" not in by


def test_build_then_assemble_roundtrip(project_docs):
    template_bytes, _, _, fields = build_template_from_examples([str(p) for p in project_docs])
    context = {
        "project_name": "Orion",
        "report_date": "2026-07-01",
        "prepared_by": "Alice Brown",
        "summary": "On track.",
        "task_status": [
            {"task": "Design", "owner": "M. Lee", "status": "Done", "due_date": "2026-07-01"},
        ],
    }
    out = assemble(template_bytes, context, fields)
    doc = Document(BytesIO(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Project Name: Orion" in texts
    assert "Report Date: 2026-07-01" in texts
    # header + exactly one rendered data row
    assert len(doc.tables[0].rows) == 2


def test_templatize_table_single_row_becomes_loop():
    # A table with only one row (no separate header) used to be left completely
    # untouched by _templatize_table (it bailed on len(rows) < 2), so its cell
    # text was never tagged AND never removed. That single row must itself
    # become the repeatable loop template.
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Item"
    t.rows[0].cells[1].text = "Owner"

    _templatize_table(t, "rows", [])

    body_xml = t._tbl.xml
    assert "{%tr for item in rows %}" in body_xml
    assert "{%tr endfor %}" in body_xml
    assert "{{ item.col1 }}" in body_xml
    assert "{{ item.col2 }}" in body_xml
    assert "Item" not in body_xml and "Owner" not in body_xml


def test_templatize_paragraph_preserves_images():
    # A paragraph that holds a logo + dynamic text: templatizing the text must
    # insert the {{ placeholder }} WITHOUT deleting the image run (regression:
    # logos used to vanish from the built template and every generated document).
    doc = Document()
    para = doc.add_paragraph("Logo: ")
    img_run = para.add_run()
    img_run._element.append(OxmlElement("w:drawing"))  # stand-in for an embedded picture
    assert _run_has_image(img_run)

    _templatize_paragraph(para, "Logo: ", "{{ company_logo }}", "")

    # Image drawing survived, and the placeholder text is present.
    assert para._p.findall(".//" + qn("w:drawing"))
    assert "{{ company_logo }}" in para.text


def _para_with_picture():
    """A paragraph carrying a minimal DrawingML picture (w:drawing > pic:cNvPr)."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run()
    drawing = OxmlElement("w:drawing")
    cnvpr = OxmlElement("pic:cNvPr")
    cnvpr.set("name", "Picture 1")
    drawing.append(cnvpr)
    run._element.append(drawing)
    return doc, para, cnvpr


def test_tag_picture_sets_replace_key():
    _doc, para, cnvpr = _para_with_picture()
    assert _paragraph_has_picture(para)
    assert _tag_picture(para, "company_logo")
    # docxtpl.replace_pic matches a picture by its cNvPr title — that's the key.
    assert cnvpr.get("title") == "company_logo"
    assert para._p.findall(f".//{{{_PIC_NS}}}cNvPr")


def test_image_bytes_decodes_inputs():
    raw = b"\x89PNG\r\n"
    import base64 as _b64

    assert _image_bytes(raw) == raw
    assert _image_bytes(_b64.b64encode(raw).decode()) == raw
    assert _image_bytes("data:image/png;base64," + _b64.b64encode(raw).decode()) == raw
    assert _image_bytes(None) is None
    assert _image_bytes("") is None


def test_build_render_context_passes_image_through():
    f = FieldDefinition(field_name="logo", label="Logo", field_type=FieldType.IMAGE, required=False)
    ctx = build_render_context([f], {"logo": "data:image/png;base64,AAAA"})
    # Image values are NOT stringified/coerced — assemble decodes + replace_pic them.
    assert ctx["logo"] == "data:image/png;base64,AAAA"


def test_neutralize_run_disarms_literal_jinja():
    # A run that already contains "{{ Client }}" must no longer read as a tag,
    # but must stay visually identical (only a zero-width space is inserted).
    out = _neutralize_run("Dear {{ Client }}, see {% if x %}note{% endif %}.")
    assert "{{" not in out and "}}" not in out
    assert "{%" not in out and "%}" not in out
    assert out.replace("​", "") == "Dear {{ Client }}, see {% if x %}note{% endif %}."


def test_build_assemble_ignores_stray_template_markers():
    # Simulate an uploaded example that is itself an ILF-style template carrying
    # its own literal "{{ ... }}" markers. The build must neutralize them so
    # docxtpl renders without a TemplateSyntaxError.
    src = Document()
    src.add_paragraph("Header: {{1Headers3}} and {% weird %} text")
    _neutralize_stray_tags(src)
    bio = BytesIO()
    src.save(bio)
    # The saved doc, treated as a template, renders cleanly with an empty context.
    out = assemble(bio.getvalue(), {}, [])
    rendered = Document(BytesIO(out))
    text = "\n".join(p.text for p in rendered.paragraphs)
    # Original markers survive visually (sans the invisible zero-width space).
    assert "{{1Headers3}}" in text.replace("​", "")
