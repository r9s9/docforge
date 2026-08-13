"""Tags-only mode: every text element becomes a field; the built template
contains only placeholders (no original prose)."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from docforge.ai_classifier import classify, derive_field_definitions
from docforge.ai_classifier.describe import describe_forced_fields
from docforge.ai_classifier.tags_only import enforce_tags_only
from docforge.assembler import assemble
from docforge.schemas.enums import ClassificationType, ElementType, FieldType, needs_field
from docforge.structure_normalizer import build_extraction
from docforge.template_builder import build_template_docx


def _make_doc(tmp_path):
    doc = Document()
    doc.add_heading("Agenda & Topics", level=1)
    doc.add_paragraph("First we discuss the overall goals of the training session.")
    doc.add_paragraph("Then we walk through the hands-on exercises together in pairs.")
    doc.add_paragraph("Finally we collect feedback from every participant.")
    doc.add_heading("Logistics", level=1)
    doc.add_paragraph("Duration: 90 minutes")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Item"
    t.rows[0].cells[1].text = "Owner"
    t.rows[1].cells[0].text = "Projector"
    t.rows[1].cells[1].text = "Alice"
    path = tmp_path / "training.docx"
    doc.save(str(path))
    return str(path)


def test_enforce_tags_only_fields_everything(tmp_path):
    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")  # heuristic engine (AI off)

    cls_by_node = {c.node_id: c for c in result.classifications}
    for e in ext.top_level_elements():
        if not (e.text or "").strip() and e.type != ElementType.TABLE:
            continue
        c = cls_by_node[e.node_id]
        if c.classification == ClassificationType.AUTO_FIELD:
            continue
        assert needs_field(c.classification), f"{e.type}: {e.text[:40]!r} stayed {c.classification}"
        assert c.field_name, f"node {e.node_id} has no field name"
        assert c.static_prefix in (None, "")
        assert c.static_suffix in (None, "")
        assert (c.description or "").strip()

    # Headings become their own text fields.
    heading = next(e for e in ext.top_level_elements() if e.type == ElementType.HEADING)
    assert cls_by_node[heading.node_id].classification == ClassificationType.DYNAMIC_TEXT

    # Consecutive body paragraphs collapse into ONE shared repeatable section.
    body = [
        e for e in ext.top_level_elements()
        if e.type == ElementType.PARAGRAPH and "we " in e.text.lower()
    ]
    names = {cls_by_node[e.node_id].field_name for e in body}
    assert len(names) == 1
    assert all(
        cls_by_node[e.node_id].classification == ClassificationType.REPEATABLE_SECTION
        for e in body
    )

    # Table is repeatable.
    table = next(e for e in ext.top_level_elements() if e.type == ElementType.TABLE)
    assert cls_by_node[table.node_id].classification == ClassificationType.REPEATABLE_TABLE

    # Field names unique per field (grouped nodes intentionally share one).
    fields = derive_field_definitions(ext, result)
    field_names = [f.field_name for f in fields]
    assert len(field_names) == len(set(field_names))
    grouped = next(f for f in fields if f.field_name in names)
    assert len(grouped.node_ids) == len(body)  # merged into one multi-node field


def test_tags_only_template_contains_no_original_text(tmp_path):
    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")
    fields = derive_field_definitions(ext, result)

    template_bytes = build_template_docx(path, result, fields)
    tpl_texts = "\n".join(p.text for p in Document(BytesIO(template_bytes)).paragraphs)

    # No original prose survives — only tags.
    for phrase in ("Agenda & Topics", "overall goals", "hands-on exercises", "90 minutes"):
        assert phrase not in tpl_texts, f"original text {phrase!r} leaked into the template"
    assert "{{" in tpl_texts  # placeholders present
    assert tpl_texts.count("{%p for") == 1  # one loop for the grouped body run

    # Generation renders 100% new text.
    ctx: dict = {}
    for f in fields:
        if f.field_type == FieldType.TABLE:
            ctx[f.field_name] = [{c.field_name: "NEWCELL" for c in f.columns}]
        elif f.classification == ClassificationType.REPEATABLE_SECTION:
            ctx[f.field_name] = ["NEWPARA1", "NEWPARA2"]
        else:
            ctx[f.field_name] = "NEWVALUE"
    out = assemble(template_bytes, ctx, fields)
    out_text = "\n".join(p.text for p in Document(BytesIO(out)).paragraphs)
    assert "NEWVALUE" in out_text and "NEWPARA1" in out_text and "NEWPARA2" in out_text
    assert "overall goals" not in out_text


def test_tags_only_leaves_nothing_a_person_wrote_untagged(tmp_path):
    """Full-template mode means it: the running header, a caption, a column
    heading and a section label all become fields. The only text left behind is
    what Word writes for itself."""
    import re as _re

    from docx.shared import Pt

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "PROJECT TITLE"
    doc.add_paragraph("REVISIONS")  # a section label
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Item"  # a column heading
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = "Projector"
    table.cell(1, 1).text = "Alice"
    table.cell(2, 0).text = "Screen"
    table.cell(2, 1).text = "Bob"
    caption = doc.add_paragraph("Table 1: Equipment list")
    caption.style = doc.styles["Caption"]
    doc.add_paragraph("The session ran for ninety minutes.")
    path = tmp_path / "everything.docx"
    doc.save(str(path))

    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")
    fields = derive_field_definitions(ext, result)
    built = Document(BytesIO(build_template_docx(path, result, fields)))

    def leftover(text: str) -> str:
        return _re.sub(r"\{\{.*?\}\}|\{%.*?%\}", "", text or "").strip()

    survivors = [leftover(p.text) for p in built.paragraphs if leftover(p.text)]
    survivors += [
        leftover(p.text)
        for t in built.tables
        for row in t.rows
        for c in row.cells
        for p in c.paragraphs
        if leftover(p.text)
    ]
    survivors += [
        leftover(p.text)
        for p in built.sections[0].header.paragraphs
        if leftover(p.text)
    ]
    assert survivors == [], f"text left untagged: {survivors}"


def test_tags_only_leaves_multi_paragraph_toc_field_untagged(tmp_path):
    # A genuine Word Table of Contents must stay a live field -- Word
    # regenerates its content from the document's real headings whenever
    # fields are updated, so tagging it as a static placeholder would either
    # get silently overwritten or ship permanently-stale headings.
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def fldchar(kind):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), kind)
        return fc

    def instr_text(text):
        it = OxmlElement("w:instrText")
        it.text = text
        return it

    doc = Document()
    doc.add_heading("Table of Contents", level=1)
    p1 = doc.add_paragraph()
    p1.add_run()._element.append(fldchar("begin"))
    p1.add_run()._element.append(instr_text(' TOC \\o "1-3" \\h \\z \\u '))
    p1.add_run()._element.append(fldchar("separate"))
    p1.add_run("Heading One\t1")
    doc.add_paragraph("Heading Two\t2")
    p3 = doc.add_paragraph("Heading Three\t3")
    p3.add_run()._element.append(fldchar("end"))
    doc.add_heading("Real Section", level=1)
    doc.add_paragraph("Genuinely fixed-vs-dynamic body content to tag.")
    path = tmp_path / "toc_doc.docx"
    doc.save(str(path))

    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")
    cls_by_node = {c.node_id: c for c in result.classifications}

    toc_nodes = [e for e in ext.top_level_elements() if e.text.startswith("Heading")]
    assert len(toc_nodes) == 3
    for e in toc_nodes:
        c = cls_by_node[e.node_id]
        assert c.classification == ClassificationType.AUTO_FIELD, (
            f"TOC entry {e.text!r} should stay AUTO_FIELD, got {c.classification}"
        )
        assert not c.field_name

    # The template must not templatize it either.
    fields = derive_field_definitions(ext, result)
    template_bytes = build_template_docx(path, result, fields)
    tpl_texts = "\n".join(p.text for p in Document(BytesIO(template_bytes)).paragraphs)
    assert "Heading One\t1" in tpl_texts or "Heading One" in tpl_texts

    # Meanwhile genuinely ordinary content in the same doc still gets tagged.
    real_body = next(
        e for e in ext.top_level_elements() if "Genuinely fixed" in e.text
    )
    assert needs_field(cls_by_node[real_body.node_id].classification)


def test_tags_only_single_row_table_gets_tagged(tmp_path):
    # A one-row "info bar" is page layout, not a list of anything: its cells
    # each become their own field, and the table keeps its shape. (It was once
    # left entirely untouched, with its text surviving into the template — the
    # invariant this still guards.)
    doc = Document()
    doc.add_heading("Details", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Trainer"
    t.rows[0].cells[1].text = "John Smith"
    path = tmp_path / "single_row.docx"
    doc.save(str(path))

    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")
    fields = derive_field_definitions(ext, result)
    template_bytes = build_template_docx(path, result, fields)
    built = Document(BytesIO(template_bytes))
    tpl_texts = built.element.body.xml
    assert "Trainer" not in tpl_texts and "John Smith" not in tpl_texts
    assert "{%tr for" not in tpl_texts, "a one-row layout table must not become a loop"
    # The row survives, and both cells are fillable.
    assert [len(tbl.rows) for tbl in built.tables] == [1]
    assert tpl_texts.count("{{ ") >= 2


def test_tags_only_unfilled_field_paragraph_is_deleted_not_blank(tmp_path):
    # A tags-only template forces EVERY paragraph into a required field; when a
    # generation doesn't supply a value for one, the paragraph should vanish
    # from the output entirely rather than leave a blank line.
    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")
    fields = derive_field_definitions(ext, result)
    template_bytes = build_template_docx(path, result, fields)

    heading_field = next(f for f in fields if f.field_name.endswith("_title"))
    ctx = {f.field_name: "" for f in fields}  # nothing filled
    # A body/table field would naturally vanish via empty-list coercion; the
    # heading is the plain scalar case this fix specifically targets.
    ctx[heading_field.field_name] = ""

    out = assemble(template_bytes, ctx, fields)
    out_paras = [p.text for p in Document(BytesIO(out)).paragraphs]
    assert not any(t == "" for t in out_paras), f"blank leftover paragraph found: {out_paras}"

    # Now fill it and confirm the paragraph reappears with the new content —
    # one more paragraph than the fully-unfilled render.
    ctx[heading_field.field_name] = "New Section Title"
    out2 = assemble(template_bytes, ctx, fields)
    out2_paras = [p.text for p in Document(BytesIO(out2)).paragraphs]
    assert "New Section Title" in out2_paras
    assert len(out2_paras) == len(out_paras) + 1


def test_describe_forced_fields_writes_ai_description(tmp_path, monkeypatch):
    from docforge.ai.client import LLMClient
    from docforge.ai.prompts import LLMFieldDescription, LLMFieldDescriptions
    from docforge.settings_store import AIConfig

    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")  # heuristic engine (AI off) -> deterministic descriptions

    heading = next(e for e in ext.top_level_elements() if e.type == ElementType.HEADING)
    cls = next(c for c in result.classifications if c.node_id == heading.node_id)
    deterministic = cls.description
    assert deterministic  # the templated fallback is already in place

    client = LLMClient(
        AIConfig(provider="openai", enabled=True, base_url="http://x", api_key="k", model="m")
    )

    def fake_complete_json(*, system, developer, user, schema, cancel_event=None):
        return LLMFieldDescriptions(
            descriptions=[
                LLMFieldDescription(node_id=heading.node_id, description="A crisp, AI-written blurb.")
            ]
        )

    monkeypatch.setattr(client, "complete_json", fake_complete_json)
    n = describe_forced_fields(client, ext, result, {heading.node_id})
    assert n == 1
    assert cls.description == "A crisp, AI-written blurb."
    assert cls.description != deterministic


def test_describe_forced_fields_applies_one_description_to_grouped_nodes(tmp_path, monkeypatch):
    from docforge.ai.client import LLMClient
    from docforge.ai.prompts import LLMFieldDescription, LLMFieldDescriptions
    from docforge.settings_store import AIConfig

    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")

    body = [
        c for c in result.classifications
        if c.classification == ClassificationType.REPEATABLE_SECTION
    ]
    assert len(body) >= 2
    name = body[0].field_name
    node_ids = {c.node_id for c in body}

    client = LLMClient(
        AIConfig(provider="openai", enabled=True, base_url="http://x", api_key="k", model="m")
    )
    calls = {"n": 0}

    def fake_complete_json(*, system, developer, user, schema, cancel_event=None):
        calls["n"] += 1
        # Only one representative node_id should have been asked about.
        rep_node = next(iter(node_ids))
        return LLMFieldDescriptions(
            descriptions=[LLMFieldDescription(node_id=rep_node, description="Shared AI description.")]
        )

    monkeypatch.setattr(client, "complete_json", fake_complete_json)
    n = describe_forced_fields(client, ext, result, node_ids)
    assert calls["n"] == 1  # one request covers the whole shared-name group
    assert n == len(body)
    assert all(c.description == "Shared AI description." for c in body if c.field_name == name)


def test_describe_forced_fields_reports_progress(tmp_path, monkeypatch):
    from docforge.ai.client import LLMClient
    from docforge.ai.prompts import LLMFieldDescriptions
    from docforge.settings_store import AIConfig

    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    result = classify(ext, None, mode="tags_only")

    forced = {c.node_id for c in result.classifications if c.field_name}
    assert forced

    client = LLMClient(
        AIConfig(provider="openai", enabled=True, base_url="http://x", api_key="k", model="m")
    )

    def fake_complete_json(*, system, developer, user, schema, cancel_event=None):
        return LLMFieldDescriptions(descriptions=[])

    monkeypatch.setattr(client, "complete_json", fake_complete_json)
    events: list[tuple[str, float, str | None]] = []
    describe_forced_fields(
        client, ext, result, forced,
        on_progress=lambda detail, frac, code=None: events.append((detail, frac, code)),
    )
    assert events, "describe pass should report at least one progress event"
    assert all(code == "describe" for _d, _f, code in events)
    # Progress must be monotonically non-decreasing and stay below the final
    # 1.0 signal, which only classify() itself emits once everything is done.
    fracs = [f for _d, f, _c in events]
    assert fracs == sorted(fracs)
    assert all(0.90 <= f < 1.0 for f in fracs)


def test_classify_emits_final_complete_signal_after_describe_pass(tmp_path, monkeypatch):
    """End-to-end: classify() itself (not classify_llm) must be the one to
    report the final 1.0/"verify" tick, and only after the tags-only describe
    pass has actually run -- otherwise the UI's progress bar looks "done"
    while the describe pass silently keeps working behind it."""
    from docforge.ai.client import LLMClient, LLMError
    from docforge.ai.prompts import LLMFieldDescriptions
    from docforge.settings_store import AIConfig

    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")

    client = LLMClient(
        AIConfig(provider="openai", enabled=True, base_url="http://x", api_key="k", model="m")
    )

    def _boom(*a, **k):
        raise LLMError("boom")

    # Heuristic-quality classify_llm path isn't under test here -- force the
    # LLM call to fail so classify() falls back to the heuristic engine, then
    # verify only describe_forced_fields + the final signal fire afterward.
    monkeypatch.setattr("docforge.ai_classifier.service.classify_llm", _boom)
    monkeypatch.setattr(client, "complete_json", lambda **kw: LLMFieldDescriptions(descriptions=[]))

    events: list[tuple[str, float, str | None]] = []
    classify(
        ext, None, client=client, mode="tags_only",
        on_progress=lambda detail, frac, code=None: events.append((detail, frac, code)),
    )
    assert events, "classify() should report progress even on the heuristic-fallback path"
    assert events[-1][1] == 1.0 and events[-1][2] == "verify"
    assert any(code == "describe" for _d, _f, code in events)
    # The final 1.0 signal is strictly the LAST thing reported.
    assert all(f < 1.0 for _d, f, _c in events[:-1])


def test_enforce_tags_only_respects_ai_named_fields(tmp_path):
    path = _make_doc(tmp_path)
    ext = build_extraction(path, "d0")
    # Simulate the AI having named a paragraph already (with a stray prefix).
    result = classify(ext, None)  # smart mode first
    body = next(
        e for e in ext.top_level_elements()
        if e.type == ElementType.PARAGRAPH and "overall goals" in e.text
    )
    c = next(c for c in result.classifications if c.node_id == body.node_id)
    c.classification = ClassificationType.DYNAMIC_TEXT
    c.field_name = "session_goals"
    c.field_type = FieldType.MULTILINE_TEXT
    c.static_prefix = "First "
    enforce_tags_only(ext, result)
    assert c.field_name == "session_goals"  # AI name kept
    assert c.static_prefix is None  # no literal text survives
