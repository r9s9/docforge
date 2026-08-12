"""A heading and its body, repeated once per item."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from docforge.ai.prompts import build_classify_prompt
from docforge.ai_classifier.fields import BLOCK_COLUMNS, derive_field_definitions
from docforge.assembler import assemble
from docforge.schemas.classification import ClassificationResult, ElementClassification
from docforge.schemas.enums import ClassificationType, FieldType
from docforge.schemas.extraction import DocumentExtraction, NormalizedElement
from docforge.schemas.template import FieldDefinition
from docforge.structure_normalizer import build_extraction
from docforge.template_builder.builder import build_template_docx


@pytest.fixture
def block_template(tmp_path):
    """A template whose one repeated group is a Heading 2 plus its paragraph."""
    doc = Document()
    doc.add_heading("Project Alpha", 2)
    doc.add_paragraph("Alpha is on track.")
    doc.add_paragraph("Prepared by the delivery team.")
    src = tmp_path / "src.docx"
    doc.save(src)

    ids = [e.node_id for e in build_extraction(str(src), "x").top_level_elements()]
    result = ClassificationResult(
        extraction_document_id="x",
        classifications=[
            ElementClassification(
                node_id=ids[0], classification=ClassificationType.REPEATABLE_BLOCK,
                field_name="projects",
            ),
            ElementClassification(
                node_id=ids[1], classification=ClassificationType.REPEATABLE_BLOCK,
                field_name="projects",
            ),
        ],
    )
    fields = [
        FieldDefinition(
            field_name="projects", field_type=FieldType.TABLE,
            classification=ClassificationType.REPEATABLE_BLOCK,
            node_ids=ids[:2], columns=BLOCK_COLUMNS(),
        )
    ]
    return build_template_docx(str(src), result, fields), fields


def _render(template: bytes, fields, rows) -> list[tuple[str, str]]:
    out = assemble(template, {"projects": rows}, fields)
    return [
        (p.style.name, p.text)
        for p in Document(BytesIO(out)).paragraphs
        if p.text.strip()
    ]


def test_the_group_becomes_a_loop_over_title_and_body(block_template):
    template, _fields = block_template
    text = [p.text for p in Document(BytesIO(template)).paragraphs]
    assert "{%p for item in projects %}" in text
    assert "{{ item.title }}" in text
    assert "{{ item.body }}" in text
    assert "{%p endfor %}" in text


def test_each_item_repeats_the_whole_group_keeping_its_styles(block_template):
    template, fields = block_template
    rendered = _render(
        template, fields,
        [
            {"title": "Project Alpha", "body": "On track."},
            {"title": "Project Beta", "body": "Delayed."},
        ],
    )
    assert rendered == [
        ("Heading 2", "Project Alpha"),
        ("Normal", "On track."),
        ("Heading 2", "Project Beta"),
        ("Normal", "Delayed."),
        ("Normal", "Prepared by the delivery team."),
    ]


def test_each_item_gets_its_own_rich_body(block_template):
    # Every item renders through the same placeholder, so without a token per
    # item they would all expand to whichever body was seen last.
    template, fields = block_template
    rendered = _render(
        template, fields,
        [
            {"title": "Alpha", "body": "Fine.\n\n- budget green\n- staffing amber"},
            {"title": "Beta", "body": "Late."},
        ],
    )
    texts = [t for _style, t in rendered]
    assert texts[:5] == ["Alpha", "Fine.", "budget green", "staffing amber", "Beta"]
    assert "Late." in texts


def test_content_outside_the_group_is_untouched(block_template):
    template, fields = block_template
    rendered = _render(template, fields, [{"title": "Only", "body": "One."}])
    assert rendered[-1] == ("Normal", "Prepared by the delivery team.")


def test_no_items_renders_nothing_for_the_group(block_template):
    template, fields = block_template
    assert _render(template, fields, []) == [("Normal", "Prepared by the delivery team.")]


def test_a_lone_paragraph_is_not_treated_as_a_titled_group(tmp_path):
    # Without a body there is no group; inventing an empty one would produce a
    # template with a blank paragraph per item.
    doc = Document()
    doc.add_heading("Only a heading", 2)
    src = tmp_path / "one.docx"
    doc.save(src)
    ids = [e.node_id for e in build_extraction(str(src), "x").top_level_elements()]
    result = ClassificationResult(
        extraction_document_id="x",
        classifications=[
            ElementClassification(
                node_id=ids[0], classification=ClassificationType.REPEATABLE_BLOCK,
                field_name="items",
            )
        ],
    )
    fields = [
        FieldDefinition(
            field_name="items", field_type=FieldType.TABLE,
            classification=ClassificationType.REPEATABLE_BLOCK,
            node_ids=[ids[0]], columns=BLOCK_COLUMNS(),
        )
    ]
    template = build_template_docx(str(src), result, fields)
    assert "{%p for item in items %}" not in [p.text for p in Document(BytesIO(template)).paragraphs]


def test_the_field_is_derived_with_title_and_body_columns():
    extraction = DocumentExtraction(
        document_id="d", filename="d.docx",
        elements=[
            NormalizedElement(node_id="n1", xpath="/", type="heading", text="Alpha"),
            NormalizedElement(node_id="n2", xpath="/", type="paragraph", text="Body."),
        ],
    )
    result = ClassificationResult(
        extraction_document_id="d",
        classifications=[
            ElementClassification(
                node_id="n1", classification=ClassificationType.REPEATABLE_BLOCK, field_name="items"
            ),
            ElementClassification(
                node_id="n2", classification=ClassificationType.REPEATABLE_BLOCK, field_name="items"
            ),
        ],
    )
    fields = derive_field_definitions(extraction, result)
    field = next(f for f in fields if f.field_name == "items")
    assert [c.field_name for c in field.columns] == ["title", "body"]
    assert field.node_ids == ["n1", "n2"]  # one field spanning the whole group


def test_the_classifier_is_only_told_about_blocks_when_enabled(monkeypatch):
    from docforge import config

    extraction = DocumentExtraction(document_id="d", filename="d.docx", elements=[])
    settings = config.get_settings()

    monkeypatch.setattr(settings, "ai_repeatable_blocks_enabled", True)
    assert "REPEATABLE_BLOCK with the SAME field_name" in build_classify_prompt(extraction, None)[1]

    monkeypatch.setattr(settings, "ai_repeatable_blocks_enabled", False)
    assert "REPEATABLE_BLOCK with the SAME field_name" not in build_classify_prompt(extraction, None)[1]
