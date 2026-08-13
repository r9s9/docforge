"""A table's layout must never depend on the data poured into it.

The failure this guards against: a cover-page table (a title block, a client
block with a logo) was treated as repeating data. Its one content row became a
loop body, so generating a document with nothing for that loop deleted the row —
taking the client's logo with it and collapsing the cover page.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.shared import Inches

from docforge.ai_classifier import derive_field_definitions
from docforge.ai_classifier.tags_only import enforce_tags_only
from docforge.assembler import assemble
from docforge.schemas.enums import ClassificationType
from docforge.structure_normalizer import build_extraction
from docforge.template_builder.builder import build_template_docx

PIC = "{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr"


def _png(path):
    """A 1x1 PNG on disk — enough for python-docx to embed a real picture."""
    path.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a"
        "0000000d49484452000000010000000108060000001f15c489"
        "0000000a49444154789c6300000002000100e527defc"
        "0000000049454e44ae426082"
    ))
    return str(path)


def _shape(data: bytes) -> dict:
    doc = Document(BytesIO(data))
    return {
        "tables": len(doc.tables),
        "rows": [len(t.rows) for t in doc.tables],
        "pictures": len(doc.element.body.findall(".//" + PIC)),
    }


@pytest.fixture
def cover_page(tmp_path):
    """A cover page like a real corporate template: layout tables, one with a logo."""
    doc = Document()
    title = doc.add_table(rows=2, cols=1)
    title.cell(0, 0).text = "PROJECT TITLE"
    title.cell(1, 0).text = "Project or document description"

    client = doc.add_table(rows=2, cols=3)
    client.cell(0, 0).text = "CLIENT NAME OR LOGO"
    client.cell(1, 0).text = "01.01.2001"
    client.cell(1, 1).text = "Document number"
    client.cell(1, 2).paragraphs[0].add_run().add_picture(_png(tmp_path / "logo.png"), width=Inches(0.4))

    doc.add_paragraph("Prepared by the delivery team.")
    src = tmp_path / "cover.docx"
    doc.save(src)
    return str(src)


def _build(src: str, *, tags_only: bool = True):
    from docforge.ai_classifier import classify

    extraction = build_extraction(src, "cover")
    result = classify(extraction)
    if tags_only:
        enforce_tags_only(extraction, result)
    fields = derive_field_definitions(extraction, result)
    return build_template_docx(src, result, fields), fields, result


def test_a_cover_page_survives_a_document_with_no_content(cover_page):
    """The regression: rendering with nothing supplied must not eat the layout."""
    original = _shape(open(cover_page, "rb").read())
    template, fields, _ = _build(cover_page)
    rendered = _shape(assemble(template, {}, fields))

    assert rendered["tables"] == original["tables"], "a table disappeared entirely"
    assert rendered["rows"] == original["rows"], "table rows were consumed by an empty loop"
    assert rendered["pictures"] == original["pictures"], "the logo was destroyed"


def test_layout_tables_are_not_turned_into_loops(cover_page):
    """Neither cover table repeats anything, so neither should become a loop."""
    _template, _fields, result = _build(cover_page)
    loops = [c for c in result.classifications if c.classification == ClassificationType.REPEATABLE_TABLE]
    assert loops == [], f"layout tables were classified as repeatable: {[c.field_name for c in loops]}"


def test_a_real_data_table_still_repeats(tmp_path):
    """The evidence that earns a loop: several rows of the same shape."""
    doc = Document()
    table = doc.add_table(rows=4, cols=3)
    for i, row in enumerate(["Item | Qty | Price", "Widget | 2 | 10", "Gadget | 1 | 25", "Cable | 5 | 3"]):
        for ci, cell in enumerate(row.split(" | ")):
            table.cell(i, ci).text = cell
    src = tmp_path / "invoice.docx"
    doc.save(src)

    _template, _fields, result = _build(str(src))
    loops = [c for c in result.classifications if c.classification == ClassificationType.REPEATABLE_TABLE]
    assert len(loops) == 1, "a table with repeating rows should still become a loop"


def test_a_logo_beside_a_caption_survives_an_empty_generation(tmp_path):
    """A picture sharing its paragraph with text must not vanish with the text."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run().add_picture(_png(tmp_path / "logo.png"), width=Inches(0.4))
    para.add_run("Project or document description")
    src = tmp_path / "logo_caption.docx"
    doc.save(src)

    template, fields, _ = _build(str(src))
    assert "Project or document description" not in Document(BytesIO(template)).element.body.xml
    assert _shape(assemble(template, {}, fields))["pictures"] == 1


def test_empty_paragraphs_do_not_become_fields(tmp_path):
    """Blank spacing paragraphs are layout, not content nobody can fill."""
    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Some real content here.")
    src = tmp_path / "spacers.docx"
    doc.save(src)

    _template, fields, _ = _build(str(src))
    assert all("spacer" not in f.field_name for f in fields), [f.field_name for f in fields]
    assert len(fields) <= 2, [f.field_name for f in fields]


def test_structural_labels_stay_boilerplate(tmp_path):
    """A word that names a part of the document is furniture, not a field."""
    doc = Document()
    for label in ("TABLE OF CONTENTS", "REVISIONS", "FIGURES", "APPENDICES"):
        doc.add_paragraph(label)
    doc.add_paragraph("The pilot ran for six weeks and met every milestone.")
    src = tmp_path / "labels.docx"
    doc.save(src)

    template, fields, _ = _build(str(src))
    body = Document(BytesIO(template)).element.body.xml
    for label in ("TABLE OF CONTENTS", "REVISIONS", "FIGURES", "APPENDICES"):
        assert label in body, f"{label} should have stayed in the template"
    assert len(fields) == 1, [f.field_name for f in fields]


def test_generated_documents_ask_word_to_rebuild_its_fields(tmp_path):
    """A table of contents caches its text; without this it describes the example."""
    import zipfile

    doc = Document()
    doc.add_paragraph("Body")
    src = tmp_path / "plain.docx"
    doc.save(src)
    template, fields, _ = _build(str(src))

    settings = zipfile.ZipFile(BytesIO(assemble(template, {}, fields))).read("word/settings.xml")
    assert b"updateFields" in settings


def test_cell_values_are_reachable_without_the_loop(cover_page):
    """A layout table's cells are ordinary fields, so their values render."""
    template, fields, _ = _build(cover_page)
    names = [f.field_name for f in fields if f.field_type.value not in ("boolean", "image")]
    values = {name: f"value-{i}" for i, name in enumerate(names)}
    text = "\n".join(p.text for p in Document(BytesIO(assemble(template, values, fields))).paragraphs)
    tables = Document(BytesIO(assemble(template, values, fields))).tables
    cell_text = " ".join(c.text for t in tables for r in t.rows for c in r.cells)
    assert "value-" in (text + cell_text), "no supplied value reached the document"
