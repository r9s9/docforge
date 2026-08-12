"""Sections that can be left out of a document without breaking its design."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from docforge.assembler import assemble
from docforge.schemas.classification import (
    ClassificationResult,
    ElementClassification,
    SectionUnderstanding,
)
from docforge.schemas.enums import ClassificationType, FieldType
from docforge.schemas.routing import RoutingResult
from docforge.schemas.template import FieldDefinition
from docforge.services.generation import _section_toggles
from docforge.structure_normalizer import build_extraction
from docforge.template_builder.builder import build_template_docx, section_toggle_name


@pytest.fixture
def two_section_template(tmp_path):
    """A report with a Summary and a Risks section, built with toggles."""
    doc = Document()
    doc.add_heading("Summary", 1)
    doc.add_paragraph("The pilot went well.")
    doc.add_heading("Risks", 1)
    doc.add_paragraph("Vendor limits are unconfirmed.")
    src = tmp_path / "src.docx"
    doc.save(src)

    ids = [e.node_id for e in build_extraction(str(src), "x").top_level_elements()]
    result = ClassificationResult(
        extraction_document_id="x",
        document_type_guess="report",
        classifications=[
            ElementClassification(
                node_id=ids[1], classification=ClassificationType.DYNAMIC_TEXT,
                field_name="overview", field_type=FieldType.MULTILINE_TEXT,
            ),
            ElementClassification(
                node_id=ids[3], classification=ClassificationType.DYNAMIC_TEXT,
                field_name="risks", field_type=FieldType.MULTILINE_TEXT,
            ),
        ],
        sections=[
            SectionUnderstanding(section_key="summary", title="Summary", field_names=["overview"]),
            SectionUnderstanding(section_key="risks", title="Risks", field_names=["risks"]),
        ],
    )
    fields = [
        FieldDefinition(
            field_name="overview", field_type=FieldType.MULTILINE_TEXT,
            classification=ClassificationType.DYNAMIC_TEXT, node_ids=[ids[1]], section_key="summary",
        ),
        FieldDefinition(
            field_name="risks", field_type=FieldType.MULTILINE_TEXT,
            classification=ClassificationType.DYNAMIC_TEXT, node_ids=[ids[3]], section_key="risks",
        ),
    ]
    return build_template_docx(str(src), result, fields), fields


def _render(template: bytes, fields, **context) -> list[str]:
    values = {"overview": "All good.", "risks": "Some risk.", **context}
    out = assemble(template, values, fields)
    return [p.text for p in Document(BytesIO(out)).paragraphs if p.text.strip()]


def test_every_section_renders_by_default(two_section_template):
    template, fields = two_section_template
    assert _render(template, fields) == ["Summary", "All good.", "Risks", "Some risk."]


def test_a_skipped_section_takes_its_heading_with_it(two_section_template):
    # A hidden section that leaves its title behind looks like a bug, not a choice.
    template, fields = two_section_template
    rendered = _render(template, fields, **{section_toggle_name("risks"): False})
    assert rendered == ["Summary", "All good."]


def test_forgetting_the_toggles_shows_everything(two_section_template):
    # An undefined variable is falsy in Jinja, so a bare "if" would make any
    # caller that skipped the toggles silently emit an empty document.
    template, fields = two_section_template
    assert len(_render(template, fields)) == 4


def test_old_templates_are_unaffected_by_toggle_variables():
    doc = Document()
    doc.add_paragraph().add_run("{{ body }}")
    buf = BytesIO()
    doc.save(buf)
    field = FieldDefinition(field_name="body", field_type=FieldType.TEXT)
    out = assemble(buf.getvalue(), {"body": "text", "_show_anything": False}, [field])
    assert [p.text for p in Document(BytesIO(out)).paragraphs] == ["text"]


# --- deciding which sections are on --------------------------------------


def _fields() -> list[FieldDefinition]:
    return [
        FieldDefinition(field_name="overview", section_key="summary"),
        FieldDefinition(field_name="risks", section_key="risks"),
    ]


def test_sections_are_on_unless_something_asked_otherwise():
    toggles = _section_toggles(_fields(), RoutingResult(template_id="t", version=1), {})
    assert toggles == {"_show_summary": True, "_show_risks": True}


def test_an_empty_field_does_not_hide_its_section():
    # Hiding must be a decision, never a side effect of thin content: a document
    # that quietly drops a section looks complete and isn't.
    routing = RoutingResult(template_id="t", version=1)
    toggles = _section_toggles(_fields(), routing, {"overview": "", "risks": ""})
    assert all(toggles.values())


def test_the_writer_can_skip_a_section():
    routing = RoutingResult(
        template_id="t", version=1,
        skipped_sections=[{"section_key": "risks", "reason": "nothing supplied"}],
    )
    toggles = _section_toggles(_fields(), routing, {})
    assert toggles["_show_risks"] is False
    assert toggles["_show_summary"] is True


def test_refine_can_skip_a_section_through_the_context():
    routing = RoutingResult(template_id="t", version=1)
    toggles = _section_toggles(_fields(), routing, {"_skipped_sections": ["summary"]})
    assert toggles["_show_summary"] is False


def test_toggle_names_are_safe_jinja_identifiers():
    assert section_toggle_name("Key Risks & Issues") == "_show_key_risks_issues"
    assert section_toggle_name("") is None
