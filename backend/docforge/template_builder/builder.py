"""template_builder — turn a representative DOCX + classifications into a
docxtpl template (spec §6, §11).

Layout is preserved by *modifying the representative document's XML in place*
rather than rebuilding it: FIXED/AUTO content is left untouched, dynamic values
are replaced with ``{{ field }}`` (keeping any static label prefix/suffix), and
repeatable tables get a ``{%tr for ... %}`` / ``{%tr endfor %}`` loop using the
verified 3-row pattern (for-row / template-row / endfor-row).

When in-place run rewriting cannot express a case cleanly, the OOXML fallback in
``ooxml_ops.py`` provides direct lxml manipulation.
"""

from __future__ import annotations

import logging
import re
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..ai_classifier import classify, derive_field_definitions
from ..ai_classifier.fields import include_field_name
from ..common.textutil import slugify_field
from ..multi_doc_differ import diff_documents, pick_representative
from ..schemas.classification import ClassificationResult
from ..schemas.enums import ClassificationType, FieldType, is_dynamic
from ..schemas.extraction import DocumentExtraction
from ..schemas.template import FieldDefinition
from ..structure_normalizer import build_extraction, walk_document

logger = logging.getLogger("docforge.template_builder")

# Jinja loop variable used inside repeatable tables/sections.
_LOOP_VAR = "item"
# The two parts of a repeated heading+body group (see ai_classifier.fields).
_BLOCK_TITLE = "title"
_BLOCK_BODY = "body"

# A field name must be a valid Jinja/Python identifier to become a placeholder
# ({{ name }}); anything else (spaces, dots, a leading digit) makes docxtpl fail
# to compile the whole template. Names are sanitized upstream, but this is the
# last-line guard so one stray name can never crash a build/preview.
_VALID_IDENT = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")


def _safe_ident(name: str | None) -> str | None:
    return name if name and _VALID_IDENT.match(name) else None


# An uploaded example may itself contain literal Jinja-like delimiters in its
# text (e.g. a Word template that already uses "{{ Client }}" markers). docxtpl
# parses the WHOLE document, so those stray tokens would crash rendering. We make
# them inert by slipping a zero-width space between the braces — invisible to the
# reader, but no longer a tag to Jinja. Our own placeholders are written AFTER
# this pass, so they stay intact.
_ZWSP = "​"
_STRAY = (("{{", "{" + _ZWSP + "{"), ("}}", "}" + _ZWSP + "}"),
          ("{%", "{" + _ZWSP + "%"), ("%}", "%" + _ZWSP + "}"))


def _neutralize_run(text: str) -> str:
    for a, b in _STRAY:
        if a in text:
            text = text.replace(a, b)
    return text


def _neutralize_stray_tags(doc) -> None:
    """Neutralize literal Jinja delimiters anywhere in the document's own text.

    Walks every paragraph the builder knows about (body, tables, headers/footers)
    and rewrites EVERY ``<w:t>`` it contains — including runs nested inside
    hyperlinks or text boxes, which ``paragraph.runs`` silently skips — so no
    stray "{{ … }}" / "{% … %}" from the example survives to confuse docxtpl.
    """
    for wn in walk_document(doc):
        if wn.kind != "paragraph":
            continue
        for t in wn.obj._p.iter(qn("w:t")):
            if t.text and ("{" in t.text or "%}" in t.text):
                t.text = _neutralize_run(t.text)


# DrawingML picture namespace — used to find a run's <pic:cNvPr> so we can tag a
# picture with a stable key. docxtpl.replace_pic() matches a picture by its
# cNvPr name/title/descr, so writing the field key into @title lets the assembler
# swap that exact picture at generation while leaving the original as the default.
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def _paragraph_has_picture(paragraph: Paragraph) -> bool:
    """True if the paragraph contains a real DrawingML picture (not just any shape)."""
    return bool(paragraph._p.findall(".//{%s}cNvPr" % _PIC_NS))


def _tag_picture(paragraph: Paragraph, key: str) -> bool:
    """Tag the first real picture in ``paragraph`` with ``key`` (its cNvPr title)."""
    cnvprs = paragraph._p.findall(".//{%s}cNvPr" % _PIC_NS)
    if not cnvprs:
        return False
    cnvprs[0].set("title", key)
    return True


def _marker_paragraph_xml(text: str):
    """A bare <w:p> carrying a docxtpl control tag (e.g. {%p if x %})."""
    p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    p.append(run)
    return p


def _insert_marker_before(paragraph: Paragraph, text: str) -> None:
    paragraph._p.addprevious(_marker_paragraph_xml(text))


def _insert_marker_after(paragraph: Paragraph, text: str) -> None:
    paragraph._p.addnext(_marker_paragraph_xml(text))


# --- run-formatting helpers -------------------------------------------------
def _run_format_at(paragraph: Paragraph, offset: int) -> dict:
    """Capture character formatting from the run covering ``offset``."""
    runs = paragraph.runs
    if not runs:
        return {}
    acc = 0
    chosen = runs[0]
    for r in runs:
        length = len(r.text or "")
        if acc <= offset < acc + length:
            chosen = r
            break
        acc += length
    else:
        chosen = runs[-1]
    f = chosen.font
    fmt: dict = {"bold": f.bold, "italic": f.italic, "underline": f.underline, "name": f.name}
    try:
        fmt["size"] = f.size
    except (AttributeError, ValueError):
        fmt["size"] = None
    try:
        fmt["color"] = f.color.rgb if (f.color is not None and f.color.rgb is not None) else None
    except (AttributeError, ValueError):
        fmt["color"] = None
    return fmt


def _run_has_image(run) -> bool:
    """True if a run carries a picture/drawing/embedded object (not just text)."""
    el = run._element
    return (
        el.find(qn("w:drawing")) is not None
        or el.find(qn("w:pict")) is not None
        or el.find(qn("w:object")) is not None
    )


def _clear_runs(paragraph: Paragraph, *, keep_images: bool = True) -> None:
    """Remove a paragraph's runs.

    By default, runs that carry an image/drawing/object are preserved: when we
    swap a paragraph's dynamic text for a ``{{ placeholder }}`` we must not delete
    its logo or picture along with the text, or images vanish from the template
    (and every document generated from it). The placeholder text is appended
    alongside the surviving image run.
    """
    for r in list(paragraph.runs):
        if keep_images and _run_has_image(r):
            continue
        r._element.getparent().remove(r._element)


def _append_run(paragraph: Paragraph, text: str, fmt: dict):
    run = paragraph.add_run(text)
    if fmt:
        if fmt.get("bold") is not None:
            run.font.bold = fmt["bold"]
        if fmt.get("italic") is not None:
            run.font.italic = fmt["italic"]
        if fmt.get("underline") is not None:
            run.font.underline = fmt["underline"]
        if fmt.get("name"):
            run.font.name = fmt["name"]
        if fmt.get("size"):
            run.font.size = fmt["size"]
        if fmt.get("color") is not None:
            try:
                run.font.color.rgb = fmt["color"]
            except (AttributeError, ValueError):
                pass
    return run


def _templatize_paragraph(paragraph: Paragraph, prefix: str, expr: str, suffix: str) -> None:
    """Rewrite a paragraph to ``prefix{{ field }}suffix`` keeping run formatting."""
    full = paragraph.text or ""
    fmt_prefix = _run_format_at(paragraph, 0)
    fmt_value = _run_format_at(paragraph, len(prefix))
    fmt_suffix = _run_format_at(paragraph, max(0, len(full) - len(suffix)))
    _clear_runs(paragraph)
    if prefix:
        _append_run(paragraph, prefix, fmt_prefix)
    _append_run(paragraph, expr, fmt_value)
    if suffix:
        _append_run(paragraph, suffix, fmt_suffix)


def _set_cell_expr(cell, expr: str) -> None:
    """Set a table cell to a single ``expr``, preserving the first run's format."""
    paras = cell.paragraphs
    if not paras:
        cell.text = expr
        return
    first = paras[0]
    fmt = _run_format_at(first, 0)
    _clear_runs(first)
    _append_run(first, expr, fmt)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _tag_row(row, text: str) -> None:
    seen: set[int] = set()
    for i, cell in enumerate(row.cells):
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        cell.text = text if i == 0 else ""


def _templatize_table(table: Table, field_name: str, columns: list) -> None:
    """Convert a data row into a repeated row driven by ``field_name``.

    Tables with >= 2 rows keep row 0 as a frozen header and loop row 1. A table
    with only ONE row has no separate header — that single row IS the loop
    template, or the whole table would otherwise be left completely untouched
    (never templatized, text never removed) since there'd be no "row 1" to pick.
    """
    rows = table.rows
    if len(rows) < 1:
        return  # no rows at all — nothing to repeat
    single_row = len(rows) == 1
    template_row = rows[0] if single_row else rows[1]

    # Rewrite each physical cell of the template row to {{ item.col }}.
    seen: set[int] = set()
    for ci, cell in enumerate(template_row.cells):
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        col = columns[ci].field_name if ci < len(columns) else f"col{ci + 1}"
        _set_cell_expr(cell, f"{{{{ {_LOOP_VAR}.{col} }}}}")

    # Drop any remaining example data rows after the template row.
    extra_start = 1 if single_row else 2
    for extra in list(table.rows)[extra_start:]:
        extra._tr.getparent().remove(extra._tr)

    # Wrap the template row with for/endfor marker rows (verified pattern).
    for_row = table.add_row()
    _tag_row(for_row, f"{{%tr for {_LOOP_VAR} in {field_name} %}}")
    endfor_row = table.add_row()
    _tag_row(endfor_row, "{%tr endfor %}")
    template_row._tr.addprevious(for_row._tr)
    template_row._tr.addnext(endfor_row._tr)


def _templatize_repeatable_block(
    paragraphs: list[Paragraph], field_name: str
) -> None:
    """Repeat a heading and its body once per item.

    The group's first paragraph carries the item's title and the second its
    body; anything after that was example prose for one item and is removed, the
    same way a grouped repeatable section keeps only its first paragraph as the
    loop template.
    """
    head, *rest = paragraphs
    _templatize_paragraph(head, "", f"{{{{ {_LOOP_VAR}.{_BLOCK_TITLE} }}}}", "")
    body = rest[0] if rest else None
    if body is not None:
        _templatize_paragraph(body, "", f"{{{{ {_LOOP_VAR}.{_BLOCK_BODY} }}}}", "")
        for extra in rest[1:]:
            _remove_paragraph(extra)
    _insert_marker_before(head, f"{{%p for {_LOOP_VAR} in {field_name} %}}")
    _insert_marker_after(body if body is not None else head, "{%p endfor %}")


def _templatize_repeatable_paragraph(paragraph: Paragraph, field_name: str) -> None:
    """Turn a paragraph into a repeated paragraph: one rendered per list item."""
    _templatize_paragraph(paragraph, "", f"{{{{ {_LOOP_VAR} }}}}", "")
    _insert_marker_before(paragraph, f"{{%p for {_LOOP_VAR} in {field_name} %}}")
    _insert_marker_after(paragraph, "{%p endfor %}")


def _remove_paragraph(paragraph: Paragraph) -> None:
    """Delete a paragraph from the document body (used for grouped sections:
    only the first paragraph of a repeatable-section field becomes the loop
    template; the rest of the original example text is removed)."""
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _wrap_optional(paragraph: Paragraph, include_name: str) -> None:
    """Wrap a paragraph so it only renders when ``include_name`` is truthy."""
    _insert_marker_before(paragraph, f"{{%p if {include_name} %}}")
    _insert_marker_after(paragraph, "{%p endif %}")


def _is_tags_only_forced(cls) -> bool:
    """Whether ``enforce_tags_only`` touched this classification.

    Marked via a rationale substring (see ``ai_classifier/tags_only.py``) so no
    schema change was needed. Used to decide whether an unfilled field's whole
    paragraph should vanish from the generated document (tags-only mode forces
    EVERY text block into a field, so many of them legitimately have no content
    for a given generation — a blank leftover line for each looks broken) versus
    a normal smart-mode dynamic field, where leaving a blank value in place is
    the existing, unsurprising behavior.
    """
    return "[tags_only]" in (cls.rationale or "")


# --- main builder -----------------------------------------------------------
def section_toggle_name(section_key: str) -> str | None:
    """The Jinja variable that decides whether a section renders at all.

    ``None`` for a key that slugifies to nothing: a bare ``_show_`` is a valid
    identifier, so every unnamed section would silently share one toggle.
    """
    slug = slugify_field(section_key, fallback="")
    return _safe_ident(f"_show_{slug}") if slug else None


def _wrap_section_spans(
    nodes: list,
    fields: list[FieldDefinition],
    sections: list,
) -> list[str]:
    """Wrap each section's paragraphs in ``{%p if _show_x %}`` … ``{%p endif %}``.

    A section is a set of fields, not a range of the document, so its span is
    inferred: from the first paragraph of its earliest field (or the heading
    immediately above it) to the paragraph before the next section starts.

    That inference can be wrong on an unusual document, and a wrong span means a
    malformed template — so every span is checked before use and simply skipped
    when it does not hold. A section that cannot be hidden is a missing feature;
    a template that will not render is a broken product.
    """
    if not sections:
        return []
    position = {wn.node_id: i for i, wn in enumerate(nodes)}
    body = [i for i, wn in enumerate(nodes) if wn.scope is None and wn.parent_node_id is None]
    if not body:
        return []

    starts: list[tuple[int, str]] = []
    for section in sections:
        key = getattr(section, "section_key", "") or ""
        member_nodes = [
            position[nid]
            for f in fields
            if f.section_key == key
            for nid in f.node_ids
            if nid in position and nid not in {n.node_id for n in nodes if n.scope is not None}
        ]
        member_nodes = [i for i in member_nodes if i in body]
        if not member_nodes:
            continue
        first = min(member_nodes)
        # Take the heading directly above the first field with the section: a
        # section hidden without its own title leaves a stranded heading.
        prev = body[body.index(first) - 1] if body.index(first) > 0 else None
        if prev is not None and nodes[prev].kind == "paragraph":
            style = (getattr(nodes[prev].obj, "style", None) and nodes[prev].obj.style.name) or ""
            if style.lower().startswith(("heading", "title")):
                first = prev
        starts.append((first, key))

    starts.sort()
    wrapped: list[str] = []
    for n, (start, key) in enumerate(starts):
        name = section_toggle_name(key)
        if not name:
            continue
        end = starts[n + 1][0] - 1 if n + 1 < len(starts) else body[-1]
        span = [i for i in body if start <= i <= end]
        # Both ends must be plain body paragraphs: a marker cannot open before a
        # table row or close inside another section's loop.
        if not span or nodes[span[0]].kind != "paragraph" or nodes[span[-1]].kind != "paragraph":
            logger.warning("skipping section toggle for %r — span is not paragraph-bounded", key)
            continue
        # Default to showing: an undefined variable renders falsy in Jinja, so a
        # bare "if" would make any caller that forgot to supply the toggles
        # silently produce an empty document. Only an explicit False hides.
        _insert_marker_before(nodes[span[0]].obj, f"{{%p if {name} is not defined or {name} %}}")
        _insert_marker_after(nodes[span[-1]].obj, "{%p endif %}")
        wrapped.append(key)
    return wrapped


def _nodes_inside_looped_tables(nodes: list, cls_by_node: dict) -> set[str]:
    """Node ids living inside a table that will be rewritten as a row loop."""
    looped = {
        wn.node_id
        for wn in nodes
        if wn.kind == "table"
        and (c := cls_by_node.get(wn.node_id)) is not None
        and c.classification == ClassificationType.REPEATABLE_TABLE
    }
    if not looped:
        return set()
    owned: set[str] = set()
    parent = {wn.node_id: wn.parent_node_id for wn in nodes}
    for node_id in parent:
        ancestor = parent.get(node_id)
        while ancestor is not None:
            if ancestor in looped:
                owned.add(node_id)
                break
            ancestor = parent.get(ancestor)
    return owned


def _templatize_blocks(nodes: list, cls_by_node: dict, fd_by_node: dict) -> set[str]:
    """Convert each repeated heading+body group into a loop; returns nodes used.

    The group is the set of consecutive paragraphs the classifier gave one
    field name — the same grouping mechanism repeatable sections use, read here
    as "a title and the body under it" rather than "one paragraph per item".
    """
    by_field: dict[str, list] = {}
    for wn in nodes:
        cls = cls_by_node.get(wn.node_id)
        if cls is None or cls.classification != ClassificationType.REPEATABLE_BLOCK:
            continue
        if wn.kind != "paragraph" or wn.scope is not None:
            continue
        fd = fd_by_node.get(wn.node_id)
        name = _safe_ident(fd.field_name if fd else cls.field_name)
        if not name:
            logger.warning("skipping repeatable block with unsafe name")
            continue
        by_field.setdefault(name, []).append(wn)

    used: set[str] = set()
    for name, group in by_field.items():
        if len(group) < 2:
            # A single paragraph is not a titled group; leave it to the ordinary
            # repeatable-section path rather than inventing an empty body.
            logger.info("repeatable block %r has no body paragraph — treating as a section", name)
            continue
        _templatize_repeatable_block([wn.obj for wn in group], name)
        used.update(wn.node_id for wn in group)
    return used


def build_template_docx(
    representative_docx_path: str,
    result: ClassificationResult,
    fields: list[FieldDefinition],
) -> bytes:
    """Produce the template.docx bytes for a representative document."""
    doc = Document(str(representative_docx_path))
    # Make any literal {{…}}/{%…%} already present in the example inert BEFORE we
    # insert our own placeholders, so docxtpl only ever parses our tags.
    _neutralize_stray_tags(doc)
    nodes = walk_document(doc)

    cls_by_node = {c.node_id: c for c in result.classifications}
    # Map node -> its VALUE field. Boolean "include_*" toggles and image fields
    # also carry the node id but are applied elsewhere (a conditional wrapper,
    # and replace_pic), so neither may become the text placeholder — a paragraph
    # holding a logo and a caption would otherwise render {{ image_1 }} as text.
    fd_by_node: dict[str, FieldDefinition] = {}
    for f in fields:
        if f.field_type in (FieldType.BOOLEAN, FieldType.IMAGE):
            continue
        for nid in f.node_ids:
            fd_by_node[nid] = f

    # Repeated heading+body groups are handled as a unit before the per-node
    # pass, since one field owns several consecutive paragraphs.
    handled_by_block = _templatize_blocks(nodes, cls_by_node, fd_by_node)

    # A looped table owns every cell inside it: its rows are generated from the
    # loop's items, so a cell that also carried its own {{ field }} would only
    # render when the loop happened to have items — and never for the value the
    # field was named after. One owner per node.
    owned_by_loop = _nodes_inside_looped_tables(nodes, cls_by_node)

    seen_section_fields: set[str] = set()
    for wn in nodes:
        if wn.node_id in handled_by_block or wn.node_id in owned_by_loop:
            continue
        cls = cls_by_node.get(wn.node_id)
        if cls is None:
            continue

        # Repeatable table.
        if cls.classification == ClassificationType.REPEATABLE_TABLE and wn.kind == "table":
            fd = fd_by_node.get(wn.node_id)
            name = _safe_ident(fd.field_name) if fd else None
            if name:
                _templatize_table(wn.obj, name, fd.columns)
            elif fd:
                logger.warning("skipping table field with unsafe name %r", fd.field_name)
            continue

        if wn.kind != "paragraph":
            continue

        para = wn.obj
        # A paragraph holding only a picture has no text to templatize — leave
        # it alone (it may be tagged below as a dynamic image field). One that
        # holds a picture *and* text still gets its text replaced: _clear_runs
        # keeps image runs, so the logo survives beside the placeholder. Without
        # this, a caption sharing a paragraph with a logo stayed literal.
        if _paragraph_has_picture(para) and not (para.text or "").strip():
            continue
        fd = fd_by_node.get(wn.node_id)
        # Compute the optional toggle name from the ORIGINAL text (before edits).
        include_name = _safe_ident(include_field_name(cls, para)) if cls.optional else None

        if cls.classification == ClassificationType.REPEATABLE_SECTION:
            name = _safe_ident(fd.field_name if fd else cls.field_name)
            if name:
                if name in seen_section_fields:
                    # Later paragraph of a grouped section — the first one already
                    # carries the loop; drop this original example text entirely.
                    _remove_paragraph(para)
                    continue
                seen_section_fields.add(name)
                _templatize_repeatable_paragraph(para, name)
        elif is_dynamic(cls.classification):
            name = _safe_ident(fd.field_name if fd else cls.field_name)
            if name:
                _templatize_paragraph(
                    para, cls.static_prefix or "", f"{{{{ {name} }}}}", cls.static_suffix or ""
                )
                if _is_tags_only_forced(cls) and not _paragraph_has_picture(para):
                    # Unfilled -> the whole paragraph disappears instead of
                    # rendering a blank line (the field's own value is both
                    # the condition and the content it guards). Never for a
                    # paragraph holding a picture: the logo beside the text is
                    # content in its own right and must not vanish with it.
                    _wrap_optional(para, name)
            elif (fd and fd.field_name) or cls.field_name:
                logger.warning(
                    "skipping field with unsafe name %r (left as fixed text)",
                    (fd.field_name if fd else cls.field_name),
                )

        if include_name:
            _wrap_optional(para, include_name)

    # Image fields: tag each underlying picture with the field key so the
    # assembler can swap it via docxtpl.replace_pic at generation. The original
    # picture stays in place — it's the default, so previews and untouched
    # ("keep original") images both render exactly as in the example.
    img_fields = [f for f in fields if f.field_type == FieldType.IMAGE]
    if img_fields:
        node_by_id = {wn.node_id: wn for wn in nodes}
        for f in img_fields:
            key = _safe_ident(f.field_name)
            if not key:
                continue
            for nid in f.node_ids:
                wn = node_by_id.get(nid)
                if wn and wn.kind == "paragraph" and _tag_picture(wn.obj, key):
                    break

    # Section toggles go in last, so the spans are measured against the nodes as
    # walked (marker paragraphs inserted here are new siblings, never renumbering
    # anything the extraction already named).
    wrapped = _wrap_section_spans(nodes, fields, result.sections)
    if wrapped:
        logger.info("section toggles added for: %s", ", ".join(wrapped))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_template_from_examples(
    example_paths: list[str],
) -> tuple[bytes, DocumentExtraction, ClassificationResult, list[FieldDefinition]]:
    """Convenience: full analyze + build directly from example file paths.

    Returns (template_docx_bytes, representative_extraction, classification, fields).
    Useful for tests, the CLI, and seed data.
    """
    extractions = [build_extraction(p, f"ex-{i}") for i, p in enumerate(example_paths)]
    diff = diff_documents(extractions) if len(extractions) >= 2 else None
    rep_i = pick_representative(extractions) if len(extractions) >= 2 else 0
    rep = extractions[rep_i]
    result = classify(rep, diff)
    fields = derive_field_definitions(rep, result)
    template_bytes = build_template_docx(str(example_paths[rep_i]), result, fields)
    return template_bytes, rep, result, fields
