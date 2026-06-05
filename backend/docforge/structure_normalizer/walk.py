"""Deterministic document walk.

Both the extractor (structure_normalizer) and the template_builder walk a DOCX
in *exactly* this order so that a ``node_id`` assigned during extraction maps
back to the same physical element when the builder templatizes it. Keep this
function pure and stable — changing the traversal order would invalidate every
previously generated template's node_id references.
"""

from __future__ import annotations

from dataclasses import dataclass

from docx.document import Document as _DocumentClass
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


@dataclass
class WalkNode:
    """One traversal stop. ``obj`` is the live python-docx Paragraph/Table."""

    node_id: str
    kind: str  # "paragraph" | "table"
    obj: object
    parent_node_id: str | None
    position_index: int
    scope: str | None  # None = document body; else "header:0" / "footer:0"
    section_index: int
    row: int | None = None  # table cell coordinates (for paragraphs inside cells)
    col: int | None = None


def iter_block_items(parent):
    """Yield Paragraphs and Tables in document order from a body or table cell.

    This is the canonical python-docx recipe — it preserves the interleaving of
    paragraphs and tables that ``.paragraphs`` / ``.tables`` would lose.
    """
    if isinstance(parent, _DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"Unsupported walk parent: {type(parent)!r}")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _has_sectpr(paragraph: Paragraph) -> bool:
    pPr = paragraph._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:sectPr")) is not None


class _IdGen:
    def __init__(self) -> None:
        self._n = 0

    def __call__(self) -> str:
        self._n += 1
        return f"n{self._n:04d}"


def _walk_table_cells(
    table: Table,
    table_nid: str,
    nodes: list[WalkNode],
    next_id: _IdGen,
    *,
    scope: str | None,
    section_index: int,
) -> None:
    """Walk the paragraphs inside each cell, skipping repeats from merged cells."""
    seen: set[int] = set()
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            key = id(cell._tc)
            if key in seen:  # merged cell already visited
                continue
            seen.add(key)
            cpos = 0
            for block in iter_block_items(cell):
                if isinstance(block, Paragraph):
                    nodes.append(
                        WalkNode(
                            next_id(), "paragraph", block, table_nid, cpos,
                            scope, section_index, r, c,
                        )
                    )
                    cpos += 1
                elif isinstance(block, Table):
                    nid = next_id()
                    nodes.append(
                        WalkNode(nid, "table", block, table_nid, cpos, scope, section_index, r, c)
                    )
                    cpos += 1
                    _walk_table_cells(block, nid, nodes, next_id, scope=scope, section_index=section_index)


def walk_document(doc) -> list[WalkNode]:
    """Return the ordered list of WalkNodes for body then headers/footers."""
    nodes: list[WalkNode] = []
    next_id = _IdGen()

    # --- body ---
    section_index = 0
    pos = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            nodes.append(WalkNode(next_id(), "paragraph", block, None, pos, None, section_index))
            pos += 1
            if _has_sectpr(block):
                section_index += 1
        elif isinstance(block, Table):
            nid = next_id()
            nodes.append(WalkNode(nid, "table", block, None, pos, None, section_index))
            pos += 1
            _walk_table_cells(block, nid, nodes, next_id, scope=None, section_index=section_index)

    # --- headers / footers ---
    # Only walk parts that are *not* inherited (linked to previous) to avoid dupes.
    for si, section in enumerate(doc.sections):
        for kind in ("header", "footer"):
            hf = getattr(section, kind, None)
            if hf is None or getattr(hf, "is_linked_to_previous", False):
                continue
            scope = f"{kind}:{si}"
            cpos = 0
            for para in hf.paragraphs:
                nodes.append(WalkNode(next_id(), "paragraph", para, None, cpos, scope, si))
                cpos += 1
            for tbl in getattr(hf, "tables", []):
                nid = next_id()
                nodes.append(WalkNode(nid, "table", tbl, None, cpos, scope, si))
                cpos += 1
                _walk_table_cells(tbl, nid, nodes, next_id, scope=scope, section_index=si)

    return nodes
