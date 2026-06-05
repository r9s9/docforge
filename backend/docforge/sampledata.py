"""Deterministic sample document generator.

Produces pairs of *similar* filled DOCX files for three document types
(project report, invoice, compliance report). Used both as test fixtures and as
seed/demo data. Builders are pure functions of ``variant`` (no clock/random) so
snapshot tests stay stable.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _add_page_number_field(paragraph) -> None:
    """Insert a Word PAGE field (an AUTO field) into ``paragraph``."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _labeled(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def _fill_table(table, headers: list[str], rows: list[list[str]]) -> None:
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for r in para.runs:
                r.bold = True
    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = val


# --- Project report ---------------------------------------------------------
_PROJECT_INTRO = (
    "This report summarizes the current status of the project, including key "
    "milestones, open risks, and next steps for the upcoming reporting period."
)


def build_project_report(variant: int) -> Document:
    data = {
        1: {
            "project": "Apollo Data Migration",
            "date": "2026-05-01",
            "author": "Jane Smith",
            "summary": "Phase 1 completed on schedule. Data validation is underway "
            "and no critical issues have been identified to date.",
            "rows": [
                ["Schema mapping", "A. Patel", "Complete", "2026-04-20"],
                ["Validation suite", "L. Chen", "In progress", "2026-05-15"],
            ],
        },
        2: {
            "project": "Helios CRM Rollout",
            "date": "2026-06-01",
            "author": "John Doe",
            "summary": "Pilot deployment reached 60% of target users. Two medium "
            "risks were escalated and mitigation plans are in place.",
            "rows": [
                ["User onboarding", "R. Gomez", "In progress", "2026-06-10"],
                ["Integration tests", "K. Müller", "Blocked", "2026-06-18"],
                ["Training material", "S. Ito", "Complete", "2026-05-29"],
            ],
        },
    }[variant]

    doc = Document()
    _title(doc, "MONTHLY PROJECT STATUS REPORT")
    _labeled(doc, "Project Name", data["project"])
    _labeled(doc, "Report Date", data["date"])
    _labeled(doc, "Prepared By", data["author"])
    doc.add_paragraph(_PROJECT_INTRO)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data["summary"])

    doc.add_heading("Task Status", level=1)
    table = doc.add_table(rows=1, cols=4)
    _fill_table(table, ["Task", "Owner", "Status", "Due Date"], data["rows"])

    doc.add_heading("Confidentiality", level=1)
    doc.add_paragraph(
        "This document is confidential and intended solely for internal use."
    )

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Confidential — Page ")
    _add_page_number_field(fp)
    return doc


# --- Invoice ----------------------------------------------------------------
_INVOICE_TERMS = (
    "Payment is due within 30 days of the invoice date. Please reference the "
    "invoice number on all remittances."
)


def build_invoice(variant: int) -> Document:
    data = {
        1: {
            "number": "INV-2026-042",
            "date": "2026-05-03",
            "bill_to": "Acme Corporation",
            "rows": [
                ["Consulting services", "10", "$150.00", "$1,500.00"],
                ["Onsite support", "2", "$400.00", "$800.00"],
            ],
            "total": "$2,300.00",
        },
        2: {
            "number": "INV-2026-0067",
            "date": "2026-06-12",
            "bill_to": "Globex LLC",
            "rows": [
                ["Software license", "5", "$220.00", "$1,100.00"],
                ["Implementation", "12", "$160.00", "$1,920.00"],
                ["Training", "3", "$300.00", "$900.00"],
            ],
            "total": "$3,920.00",
        },
    }[variant]

    doc = Document()
    _title(doc, "INVOICE")
    _labeled(doc, "Invoice Number", data["number"])
    _labeled(doc, "Invoice Date", data["date"])
    _labeled(doc, "Bill To", data["bill_to"])

    doc.add_heading("Line Items", level=1)
    table = doc.add_table(rows=1, cols=4)
    _fill_table(table, ["Description", "Qty", "Unit Price", "Amount"], data["rows"])

    _labeled(doc, "Total Due", data["total"])

    doc.add_heading("Terms", level=1)
    doc.add_paragraph(_INVOICE_TERMS)
    return doc


# --- Compliance report ------------------------------------------------------
_COMPLIANCE_SCOPE = (
    "This compliance report documents the results of the audit conducted against "
    "the applicable control framework. Findings are listed in the table below."
)


def build_compliance_report(variant: int) -> Document:
    data = {
        1: {
            "standard": "ISO 27001:2022",
            "audit_date": "2026-04-15",
            "auditor": "Maria Lopez",
            "rows": [
                ["A.5.1", "Information security policies", "Pass", "Low"],
                ["A.8.2", "Privileged access rights", "Finding", "High"],
            ],
        },
        2: {
            "standard": "SOC 2 Type II",
            "audit_date": "2026-05-20",
            "auditor": "David Okafor",
            "rows": [
                ["CC6.1", "Logical access controls", "Pass", "Low"],
                ["CC7.2", "Security monitoring", "Finding", "Medium"],
                ["CC8.1", "Change management", "Pass", "Low"],
            ],
        },
    }[variant]

    doc = Document()
    _title(doc, "COMPLIANCE AUDIT REPORT")
    _labeled(doc, "Standard", data["standard"])
    _labeled(doc, "Audit Date", data["audit_date"])
    _labeled(doc, "Lead Auditor", data["auditor"])
    doc.add_paragraph(_COMPLIANCE_SCOPE)

    doc.add_heading("Findings", level=1)
    table = doc.add_table(rows=1, cols=4)
    _fill_table(table, ["Control", "Description", "Result", "Severity"], data["rows"])

    doc.add_heading("Statement", level=1)
    doc.add_paragraph(
        "The auditor attests that the assessment was performed in accordance "
        "with the stated methodology and professional standards."
    )
    return doc


# --- Service agreement with OPTIONAL content (for robustness/optional tests) --
# Variant 1 includes a special clause paragraph that variant 2 omits -> the
# differ should mark it optional, the builder should wrap it in {%p if %}.
def build_service_agreement(variant: int) -> Document:
    doc = Document()
    _title(doc, "SERVICE AGREEMENT")
    _labeled(doc, "Reference", f"SA-2026-{variant:03d}")
    doc.add_paragraph(
        "This agreement sets out the standard terms of service between the parties."
    )
    if variant == 1:
        doc.add_paragraph(
            "Special Clause: Expedited 24/7 support is included for this client."
        )
    doc.add_heading("Signatures", level=1)
    doc.add_paragraph("Signed by both parties on the date written above.")
    return doc


# --- registry ---------------------------------------------------------------
SAMPLE_SETS: dict[str, list] = {
    "project_report": [build_project_report, build_project_report],
    "invoice": [build_invoice, build_invoice],
    "compliance_report": [build_compliance_report, build_compliance_report],
}

_BUILDERS = {
    "project_report": build_project_report,
    "invoice": build_invoice,
    "compliance_report": build_compliance_report,
}


def write_sample(kind: str, variant: int, path: str | Path) -> Path:
    """Build one sample document and save it to ``path``."""
    builder = _BUILDERS[kind]
    doc = builder(variant)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_all(out_dir: str | Path) -> dict[str, list[Path]]:
    """Write all three document types (2 variants each) to ``out_dir``.

    Returns {kind: [path_variant1, path_variant2]}.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    result: dict[str, list[Path]] = {}
    for kind in _BUILDERS:
        paths = []
        for variant in (1, 2):
            p = out_dir / f"{kind}_{variant}.docx"
            write_sample(kind, variant, p)
            paths.append(p)
        result[kind] = paths
    return result
