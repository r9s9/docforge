"""In-place compliance fix.

Repairs a document's FIXED (boilerplate) text so it matches the template exactly,
while leaving everything else — dynamic field values, extra content, formatting —
untouched. We align the upload to the template's representative structure, then
for each FIXED element that was changed or removed we patch the corresponding
paragraph in the *original* uploaded DOCX (located via the deterministic walk).
"""

from __future__ import annotations

import os
import tempfile
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from ..common.textutil import similarity
from ..config import Settings, get_settings
from ..db.models import Template
from ..multi_doc_differ import align_to_representative
from ..schemas.enums import ClassificationType
from ..schemas.extraction import DocumentExtraction
from ..structure_normalizer import build_extraction, walk_document
from ..template_registry import TemplateRegistry

_FIXED_MATCH_THRESHOLD = 0.9


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paragraph.add_run(text)


def _insert_after(anchor: Paragraph, text: str) -> Paragraph:
    """Insert a new paragraph carrying ``text`` immediately after ``anchor``."""
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    para.add_run(text)
    return para


def fix_document(
    db: Session,
    template: Template,
    *,
    filename: str,
    data: bytes,
    version: int | None = None,
    settings: Settings | None = None,
    registry: TemplateRegistry | None = None,
) -> tuple[bytes, int]:
    """Return (corrected_docx_bytes, number_of_fixes_applied)."""
    settings = settings or get_settings()
    registry = registry or TemplateRegistry(settings.templates_dir)
    version = version or template.latest_version

    rep_raw = registry.load_representative(template.id, version)
    if rep_raw is None:
        raise ValueError(
            "This template version has no stored representative structure; "
            "re-publish the template to enable fixing."
        )
    rep = DocumentExtraction.model_validate(rep_raw)
    intelligence = registry.load_intelligence(template.id, version)
    cls_by_node = {c.node_id: c for c in intelligence.classifications}

    # Extract for alignment, and load the live python-docx for patching.
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    try:
        tmp.write(data)
        tmp.close()
        extraction = build_extraction(tmp.name, document_id="fix", filename=filename)
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    pydoc = Document(BytesIO(data))
    aligned = align_to_representative(rep, extraction)
    walk = {n.node_id: n for n in walk_document(pydoc)}

    n_fixed = 0
    last_anchor: Paragraph | None = None  # nearest preceding mapped paragraph (for inserts)

    for node in rep.top_level_elements():
        c = cls_by_node.get(node.node_id)
        doc_el = aligned.get(node.node_id)

        # Track the last doc paragraph that maps to a template element, so a
        # missing-fixed paragraph can be inserted in roughly the right place.
        if doc_el is not None:
            wn = walk.get(doc_el.node_id)
            if wn is not None and wn.kind == "paragraph":
                last_anchor = wn.obj

        if c is None or c.classification != ClassificationType.FIXED:
            continue
        expected = node.text.strip()
        if not expected:
            continue

        if doc_el is None:
            # Boilerplate paragraph removed from the upload — re-insert it.
            if last_anchor is not None:
                last_anchor = _insert_after(last_anchor, expected)
                n_fixed += 1
            continue

        wn = walk.get(doc_el.node_id)
        if wn is None or wn.kind != "paragraph":
            continue
        para = wn.obj
        if similarity((para.text or "").strip(), expected) < _FIXED_MATCH_THRESHOLD:
            _set_paragraph_text(para, expected)
            n_fixed += 1

    out = BytesIO()
    pydoc.save(out)
    return out.getvalue(), n_fixed
